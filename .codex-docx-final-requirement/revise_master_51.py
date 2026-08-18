from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"D:\ฝึกงาน\Requirement\ระบบอนุมัติใบลาออนไลน์ Final version.docx")
OUTPUT = ROOT / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"
WORKING = ROOT / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.master-working.docx"


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_paragraph(document, exact: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact}")


def replace_exact(document, old: str, new: str) -> None:
    set_text(find_paragraph(document, old), new)


def clone_paragraph_before(template, anchor, text: str):
    node = deepcopy(template._p)
    anchor._p.addprevious(node)
    paragraph = template._parent.paragraphs[0]
    for candidate in template._parent.paragraphs:
        if candidate._p is node:
            paragraph = candidate
            break
    set_text(paragraph, text)
    return paragraph


def insert_block(document, anchor_text: str, items: list[tuple[str, str]]) -> None:
    anchor = find_paragraph(document, anchor_text)
    normal_template = find_paragraph(document, "FR-01 Authentication")
    bullet_template = find_paragraph(document, "ผู้ใช้ Logout ได้")
    body_template = find_paragraph(document, "ระบบต้องรองรับการเข้าสู่ระบบ ออกจากระบบ และเปลี่ยนรหัสผ่าน")
    detail_template = find_paragraph(document, "รายละเอียด")
    templates = {
        "heading": normal_template,
        "bullet": bullet_template,
        "body": body_template,
        "label": detail_template,
    }
    for kind, text in items:
        clone_paragraph_before(templates[kind], anchor, text)


def replace_paragraph_range(document, start_text: str, end_text: str, new_items: list[str]) -> None:
    start = find_paragraph(document, start_text)
    end = find_paragraph(document, end_text)
    template = find_paragraph(document, "ผู้ใช้ Logout ได้")
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node
    for text in new_items:
        clone_paragraph_before(template, end, text)


def set_cell_text(cell, value: str, bold: bool = False, center: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = str(value)
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def rewrite_table(table, headers: list[str], rows: list[list[str]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows[0].cells) < len(headers):
        table.rows[0]._tr.append(deepcopy(table.rows[0].cells[-1]._tc))
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, center=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value, center=index == 0 and len(headers) <= 3)


def append_table_rows(table, rows: list[list[str]]) -> None:
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, center=index == 0)


def table_after_heading(document, heading: str):
    paragraph = find_paragraph(document, heading)
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            for table in document.tables:
                if table._tbl is node:
                    return table
        node = node.getnext()
    raise ValueError(f"Table after heading not found: {heading}")


def add_table_before(document, anchor_text: str, headers: list[str], rows: list[list[str]], widths: list[int]):
    anchor = find_paragraph(document, anchor_text)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True, center=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, center=index == 0)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
    anchor._p.addprevious(table._tbl)
    return table


def global_text_replacements(document) -> None:
    replacements = {
        "ระบบมี 4 Role หลัก": "ระบบมี 4 Role หลัก ได้แก่ Employee, Supervisor, HR และ Admin",
        "ระบบใช้สถานะหลัก 5 สถานะ": "ระบบใช้สถานะหลัก 5 สถานะสำหรับ Workflow การลาแบบ Supervisor อนุมัติขั้นเดียว",
        "ระบบ MVP ยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว": "ระบบเวอร์ชันปัจจุบันยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว",
        "หมายเหตุ: ระบบ MVP ยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว ดังนั้นยังไม่มีการคืนวันลาในเวอร์ชันแรก": "หมายเหตุ: ระบบเวอร์ชันปัจจุบันยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว จึงยังไม่มี Workflow คืนสิทธิ์วันลา",
        "หมายเหตุ: ระบบเวอร์ชันปัจจุบันยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว ดังนั้นยังไม่มีการคืนวันลาในเวอร์ชันแรก": "หมายเหตุ: ระบบเวอร์ชันปัจจุบันยังไม่รองรับการยกเลิกคำขอที่ Approved แล้ว จึงยังไม่มี Workflow คืนสิทธิ์วันลา",
        "Role เป็น master data ของระบบ ใน MVP ไม่เปิดให้เพิ่ม Role ใหม่ผ่านหน้าจอ": "Role เป็น Master Data ของระบบเวอร์ชันปัจจุบัน และไม่เปิดให้เพิ่ม Role ใหม่ผ่านหน้าจอ",
        "หมายเหตุที่แก้จากเดิม": "หมายเหตุ",
        "ระบบใช้ 13 ตารางหลัก": "ระบบใช้ 15 ตารางหลัก",
        "14. ข้อจำกัดของระบบเวอร์ชัน MVP": "14. ข้อจำกัดของระบบเวอร์ชันปัจจุบัน",
        "ระบบไม่รองรับ Email Notification": "ระบบไม่ใช้ Email เป็นช่องทางแจ้งเตือน Workflow การลา โดย Email ใช้เฉพาะส่ง OTP สำหรับ Forgot Password",
        "ระบบไม่รองรับ Advanced Dashboard": "Dashboard ปัจจุบันแสดงข้อมูลสรุปตาม Role แต่ยังไม่รองรับการวิเคราะห์ขั้นสูงหลายมิติ",
        "Then ระบบต้องไม่ให้ส่งคำขอ": "Then ระบบต้องไม่ให้ส่งคำขอ",
        "ตาราง notifications ใช้สำหรับจัดเก็บข้อมูลการแจ้งเตือนภายในระบบ เพื่อแจ้งผู้ใช้งานที่เกี่ยวข้องเมื่อมีเหตุการณ์สำคัญ เช่น มีคำขอรออนุมัติ คำขอได้รับการอนุมัติ คำขอถูกปฏิเสธ หรือคำขอถูกยกเลิก": "ตาราง notifications ใช้สำหรับจัดเก็บ In-app Notification ของผู้ใช้ โดย Workflow ปัจจุบันสร้างรายการเมื่อ Submit, Approve หรือ Reject",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            set_text(paragraph, new_text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    new_text = text
                    if new_text != text:
                        set_text(paragraph, new_text)


def build() -> None:
    shutil.copy2(SOURCE, WORKING)
    document = Document(WORKING)
    global_text_replacements(document)

    replace_exact(document, "Online Leave Approval System", "Online Leave Approval System\nFinal Requirement")
    replace_exact(document, "2.2 ฟังก์ชันที่ไม่อยู่ในขอบเขต MVP", "2.2 ขอบเขตในอนาคต (Future Scope)")
    replace_exact(document, "2.1 ฟังก์ชันที่อยู่ในขอบเขต MVP", "2.1 ฟังก์ชันที่อยู่ในขอบเขตระบบเวอร์ชันปัจจุบัน")

    scope_table = table_after_heading(document, "2.1 ฟังก์ชันที่อยู่ในขอบเขตระบบเวอร์ชันปัจจุบัน")
    append_table_rows(scope_table, [
        ["Password Recovery", "Forgot Password ด้วย OTP ทาง Email และ Reset Token แบบใช้ครั้งเดียว"],
        ["Profile Management", "ผู้ใช้ทุก Role ดูและแก้ไข Profile รวมถึงรูปโปรไฟล์ของตนเอง"],
        ["Role Dashboard", "Dashboard แสดงข้อมูลสรุปและทางลัดตามสิทธิ์ของแต่ละ Role"],
    ])

    future_table = table_after_heading(document, "2.2 ขอบเขตในอนาคต (Future Scope)")
    rewrite_table(future_table, ["Feature", "รายละเอียด", "สถานะ"], [
        ["Multi-role User", "ผู้ใช้หนึ่งคนมีหลาย Role", "Future Scope"],
        ["Half-day Leave", "ลาครึ่งวันช่วงเช้าหรือบ่าย", "Future Scope"],
        ["Cross-year Leave Request", "ยื่นคำขอลาคร่อมปี", "Future Scope"],
        ["Cancel After Approval", "ยกเลิกหลังอนุมัติพร้อมคืนสิทธิ์วันลา", "Future Scope"],
        ["Backup Assignee", "ระบุผู้รับผิดชอบงานแทนระหว่างลา", "Future Scope"],
        ["Multi-step Approval", "Department Manager, HR หรือ Final Approver อนุมัติหลายระดับ", "Future Scope"],
        ["Backup / Delegate Approver", "มอบหมายผู้อนุมัติแทนตามช่วงเวลา", "Future Scope"],
        ["Push / Real-time Notification", "แจ้งเตือนทันทีหรือส่งเข้ามือถือ", "Future Scope"],
        ["Native Mobile Application", "แอปพลิเคชันสำหรับอุปกรณ์เคลื่อนที่", "Future Scope"],
        ["PDF Export", "ส่งออกรายงานเป็น PDF", "Future Scope"],
        ["Advanced Analytics", "Dashboard และรายงานเชิงวิเคราะห์หลายมิติ", "Future Scope"],
        ["Backup / Restore", "สำรองและกู้คืนระบบระดับโครงสร้างพื้นฐาน", "Future Scope"],
        ["Large-scale Load Test", "ทดสอบภาระงานผู้ใช้พร้อมกันจำนวนมาก", "Future Scope"],
    ])

    permission_table = table_after_heading(document, "5.5 Permission Matrix")
    append_table_rows(permission_table, [
        ["Forgot Password / OTP / Reset Password", "✓", "✓", "✓", "✓"],
        ["แก้ไข Profile และรูปโปรไฟล์ของตนเอง", "✓", "✓", "✓", "✓"],
        ["ดู/อ่าน/อ่านทั้งหมด/ลบ Notification ของตนเอง", "✓", "✓", "✓", "✓"],
        ["Reset Password และออก Temporary Password ให้ผู้ใช้", "✗", "✗", "✗", "✓"],
        ["เรียกใช้ API Employee และ Master Data ของ HR", "✗", "✗", "✓", "✓"],
        ["ดาวน์โหลด Attachment ตาม Data Scope", "✓", "✓", "✓", "✓"],
    ])

    replace_exact(document, "เมื่อคำขอถูกยกเลิก", "Notification Event ของ Workflow ปัจจุบันมีเฉพาะ Submit, Approve และ Reject; การ Cancel ไม่สร้าง Notification")
    replace_exact(document, "Admin เห็นเฉพาะ metadata และไม่สามารถ download/open file ได้", "Admin สามารถเปิดหรือดาวน์โหลดไฟล์แนบตามสิทธิ์ของ Backend ปัจจุบัน")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == "Admin เห็นเฉพาะ metadata ของเอกสารแนบ":
                        set_text(paragraph, "Admin เข้าถึงเอกสารแนบได้ตามสิทธิ์ของ Backend ปัจจุบัน")
    replace_exact(document, "Admin เห็นเฉพาะข้อมูลเชิงระบบและ metadata ของเอกสาร ไม่ควรเห็นเหตุผลการลาหรือเอกสารส่วนตัวของพนักงานอื่น", "Admin เข้าถึงข้อมูลตาม Endpoint และ Data Scope ที่ Backend อนุญาต โดยต้องผ่าน Authentication และ Authorization")
    notif_table = table_after_heading(document, "10.12 Notification")
    rewrite_table(notif_table, ["Method", "Endpoint", "Permission"], [
        ["GET", "/api/notifications", "ทุก Role เฉพาะของตนเอง"],
        ["PATCH", "/api/notifications/{notificationId}/read", "ทุก Role เฉพาะของตนเอง"],
        ["PATCH", "/api/notifications/read-all", "ทุก Role เฉพาะของตนเอง"],
        ["DELETE", "/api/notifications/{notificationId}", "ทุก Role เฉพาะของตนเอง"],
    ])

    insert_block(document, "8. Business Rules และ Non-Functional Requirements", [
        ("heading", "FR-17 Temporary Password และ Password Policy"),
        ("body", "ระบบต้องรองรับ Temporary Password สำหรับบัญชีที่ Admin สร้างหรือ Reset Password และบังคับให้ผู้ใช้เปลี่ยนรหัสผ่านก่อนใช้งานฟังก์ชันธุรกิจ"),
        ("label", "รายละเอียด"),
        ("bullet", "Admin สร้างบัญชีหรือ Reset Password แล้วระบบสร้าง Temporary Password แบบสุ่ม"),
        ("bullet", "ระบบจัดเก็บเฉพาะ Password Hash และกำหนด must_change_password = true"),
        ("bullet", "ผู้ใช้ที่ยังไม่เปลี่ยน Temporary Password เข้าถึงได้เฉพาะ Session และ Change Password"),
        ("bullet", "รหัสผ่านใหม่ต้องยาวอย่างน้อย 10 ตัว มีตัวพิมพ์ใหญ่ ตัวพิมพ์เล็ก ตัวเลข และอักขระพิเศษ โดยห้ามมีช่องว่าง"),
        ("bullet", "รหัสผ่านใหม่ต้องต่างจากรหัสผ่านปัจจุบัน Username และ Email"),
        ("bullet", "เมื่อเปลี่ยนรหัสผ่านสำเร็จ ระบบล้าง must_change_password และเพิ่ม token_version เพื่อยกเลิก Session เดิม"),
        ("heading", "FR-18 Forgot Password, OTP และ Reset Token"),
        ("body", "ผู้ใช้ทุก Role สามารถขอรีเซ็ตรหัสผ่านด้วย Username หรือ Email ที่ผูกกับบัญชี Active ได้"),
        ("label", "รายละเอียด"),
        ("bullet", "ระบบส่ง OTP 6 หลักไปยัง Email ที่ลงทะเบียน โดยไม่เปิดเผยว่ามีบัญชีหรือไม่"),
        ("bullet", "ระบบจำกัดอัตราการขอ OTP, ระยะเวลาขอซ้ำ และจำนวนครั้งที่กรอกผิด"),
        ("bullet", "OTP ต้องมีวันหมดอายุ ใช้ได้ครั้งเดียว และ OTP ใหม่ต้องทำให้ OTP เดิมใช้ไม่ได้"),
        ("bullet", "เมื่อยืนยัน OTP สำเร็จ ระบบออก Reset Token แบบใช้ครั้งเดียวและมีวันหมดอายุ"),
        ("bullet", "การ Reset Password ต้องตรวจ Password Policy, ทำให้ OTP/Token ถูกใช้แล้ว และยกเลิก Token อื่นของผู้ใช้"),
        ("heading", "FR-19 Profile Management"),
        ("body", "ผู้ใช้ทุก Role สามารถดูและแก้ไขข้อมูล Profile ของตนเองได้ตามฟิลด์ที่ระบบอนุญาต"),
        ("label", "รายละเอียด"),
        ("bullet", "ผู้ใช้แก้ไขชื่อ นามสกุล Email เบอร์โทรศัพท์ และรูปโปรไฟล์ของตนเองได้"),
        ("bullet", "รูปโปรไฟล์รองรับ JPEG, PNG หรือ WebP จำนวน 1 ไฟล์ ขนาดไม่เกิน 2 MB"),
        ("bullet", "ผู้ใช้สามารถลบรูปเดิมได้ และระบบต้องไม่อนุญาตให้แก้ Employee ID, Role, Department หรือ Position ผ่านหน้า Profile"),
        ("heading", "FR-20 Notification Management"),
        ("body", "ผู้ใช้ทุก Role จัดการ In-app Notification ที่เป็นของตนเองได้"),
        ("label", "รายละเอียด"),
        ("bullet", "ระบบแสดงจำนวน Unread และเรียงรายการตามเวลาล่าสุด"),
        ("bullet", "ผู้ใช้ Mark Read, Mark All Read และ Delete Notification ของตนเองได้"),
        ("bullet", "Notification ที่มี Leave Request อ้างอิงสามารถเปิด Deep Link ไปยังรายละเอียดที่ผู้ใช้มีสิทธิ์เข้าถึง"),
        ("bullet", "Submit แจ้ง Supervisor; Approve หรือ Reject แจ้งเจ้าของคำขอ"),
        ("bullet", "HR และ Admin มีหน้า Notification ได้ แต่ Workflow ปัจจุบันไม่สร้าง Event เฉพาะให้สอง Role นี้"),
        ("heading", "FR-21 Role Dashboard"),
        ("body", "ระบบต้องแสดง Dashboard ตาม Role โดยใช้ข้อมูลจริงจาก Backend และจำกัดข้อมูลตาม Ownership หรือ Data Scope"),
        ("label", "รายละเอียด"),
        ("bullet", "ทุก Role เห็นข้อมูลคำขอและ Leave Balance ของตนเอง"),
        ("bullet", "Supervisor เห็นจำนวน Pending และรายการล่าสุดของ Direct Team"),
        ("bullet", "HR เห็นข้อมูลสรุปพนักงาน ประเภทการลา วันหยุด และคำขอระดับองค์กรตามสิทธิ์"),
        ("bullet", "Admin เห็นข้อมูลสรุปบัญชีผู้ใช้ โครงสร้างองค์กร และ Audit Log ตามสิทธิ์"),
    ])

    br_table = table_after_heading(document, "8.1 Business Rules")
    append_table_rows(br_table, [
        ["BR-32", "Reason ต้องมีความยาว 5-500 ตัวอักษร"],
        ["BR-33", "จำนวนวันลาต้องอยู่ระหว่าง minimum_days และ maximum_days_per_request ของ Leave Type"],
        ["BR-34", "ไฟล์แนบคำขอลารองรับสูงสุด 5 ไฟล์ต่อการอัปโหลด และตรวจ MIME Type"],
        ["BR-35", "Notification อ่าน แก้สถานะ หรือลบได้เฉพาะผู้ใช้เจ้าของข้อมูล"],
        ["BR-36", "Temporary Password ต้องบังคับเปลี่ยนก่อนเข้าถึงฟังก์ชันธุรกิจ"],
        ["BR-37", "OTP และ Reset Token ต้องมีวันหมดอายุ ใช้ได้ครั้งเดียว และถูกจัดเก็บในรูปแบบ Hash"],
        ["BR-38", "Username ต้องไม่ซ้ำ และพนักงานหนึ่งคนมี User Account ได้หนึ่งบัญชี"],
        ["BR-39", "ชื่อ Department, Position, Leave Type Code และ Entitlement ต่อพนักงาน/ประเภทลา/ปีต้องไม่ซ้ำ"],
        ["BR-40", "ระบบบันทึกเฉพาะ Event สำคัญที่กำหนดใน Audit Log ไม่รับประกันการบันทึกทุก Action"],
    ])

    # Database tables and fields added by the current migrations.
    users = table_after_heading(document, "9.3 ตาราง users")
    append_table_rows(users, [
        ["password_changed_at", "DATETIME", "NULL", "เวลาที่เปลี่ยนรหัสผ่านล่าสุด"],
        ["must_change_password", "BOOLEAN", "", "บังคับเปลี่ยน Temporary Password"],
        ["token_version", "INT UNSIGNED", "", "รุ่น Token สำหรับยกเลิก Session เดิม"],
    ])
    employees = table_after_heading(document, "9.4 ตาราง employees")
    append_table_rows(employees, [["profile_image_url", "VARCHAR(255)", "NULL", "ตำแหน่งไฟล์รูปโปรไฟล์"]])
    leave_types = table_after_heading(document, "9.7 ตาราง leave_types")
    append_table_rows(leave_types, [
        ["leave_type_code", "VARCHAR(10)", "UNIQUE", "รหัสประเภทการลา"],
        ["description", "VARCHAR(300)", "NULL", "รายละเอียดประเภทการลา"],
        ["minimum_days", "DECIMAL(5,2)", "", "จำนวนวันลาขั้นต่ำต่อคำขอ"],
        ["maximum_days_per_request", "DECIMAL(5,2)", "", "จำนวนวันลาสูงสุดต่อคำขอ"],
    ])
    holidays = table_after_heading(document, "9.12 ตาราง holidays")
    append_table_rows(holidays, [
        ["holiday_type", "VARCHAR(50)", "", "ประเภทวันหยุด"],
        ["description", "VARCHAR(300)", "NULL", "รายละเอียดวันหยุด"],
    ])

    replace_exact(document, "leave_attachments", "leave_request_attachments")
    replace_exact(document, "9.11 ตาราง leave_attachments", "9.11 ตาราง leave_request_attachments")
    replace_exact(document, "ตาราง leave_attachments ใช้สำหรับจัดเก็บข้อมูลเอกสารแนบประกอบคำขอลา เช่น ใบรับรองแพทย์ หรือเอกสารประกอบอื่น ๆ", "ตาราง leave_request_attachments ใช้สำหรับจัดเก็บ Metadata ของเอกสารแนบประกอบคำขอลา")
    attachments = table_after_heading(document, "9.11 ตาราง leave_request_attachments")
    rewrite_table(attachments, ["Field", "Type", "Key", "Description"], [
        ["attachment_id", "BIGINT UNSIGNED", "PK", "รหัสเอกสารแนบ"],
        ["leave_request_id", "INT UNSIGNED", "FK", "อ้างอิงคำขอลาและลบตามคำขอ"],
        ["original_name", "VARCHAR(255)", "", "ชื่อไฟล์เดิม"],
        ["stored_name", "VARCHAR(255)", "", "ชื่อไฟล์ที่จัดเก็บ"],
        ["mime_type", "VARCHAR(100)", "", "MIME Type ที่ตรวจสอบแล้ว"],
        ["file_size", "BIGINT UNSIGNED", "", "ขนาดไฟล์หน่วย Byte"],
        ["created_at", "DATETIME", "", "วันที่และเวลาที่อัปโหลด"],
    ])

    insert_block(document, "10. API Specification", [
        ("heading", "9.15 ตาราง password_reset_otps"),
        ("body", "ตาราง password_reset_otps จัดเก็บ OTP แบบ Hash พร้อมสถานะการยืนยัน การใช้งาน การยกเลิก จำนวนครั้งที่ลอง และข้อมูลควบคุมอัตราการขอ"),
    ])
    add_table_before(document, "10. API Specification", ["Field", "Type", "Key", "Description"], [
        ["id", "BIGINT UNSIGNED", "PK", "รหัสรายการ OTP"],
        ["user_id", "INT UNSIGNED", "FK", "บัญชีผู้ขอ Reset Password"],
        ["otp_hash", "CHAR(64)", "", "ค่า OTP ที่ Hash แล้ว"],
        ["expires_at", "DATETIME", "", "เวลาหมดอายุ"],
        ["verified_at / used_at", "DATETIME", "NULL", "เวลายืนยันและใช้งาน"],
        ["invalidated_at", "DATETIME", "NULL", "เวลาที่ทำให้ใช้ไม่ได้"],
        ["attempt_count / resend_count", "TINYINT", "", "จำนวนครั้งที่ลองและขอซ้ำ"],
        ["requested_ip / requested_user_agent", "VARCHAR", "NULL", "ข้อมูลคำขอเพื่อความปลอดภัย"],
    ], [2337, 2337, 2338, 2338])
    insert_block(document, "10. API Specification", [
        ("heading", "9.16 ตาราง password_reset_tokens"),
        ("body", "ตาราง password_reset_tokens จัดเก็บ Reset Token แบบ Hash ที่ผูกกับ OTP ซึ่งผ่านการยืนยันแล้วและใช้ได้ครั้งเดียว"),
    ])
    add_table_before(document, "10. API Specification", ["Field", "Type", "Key", "Description"], [
        ["id", "BIGINT UNSIGNED", "PK", "รหัส Reset Token"],
        ["user_id", "INT UNSIGNED", "FK", "บัญชีผู้ใช้"],
        ["token_hash", "CHAR(64)", "UNIQUE", "Reset Token ที่ Hash แล้ว"],
        ["otp_id", "BIGINT UNSIGNED", "FK", "OTP ที่ผ่านการยืนยัน"],
        ["expires_at", "DATETIME", "", "เวลาหมดอายุ"],
        ["used_at / invalidated_at", "DATETIME", "NULL", "สถานะใช้แล้วหรือถูกยกเลิก"],
        ["created_at", "DATETIME", "", "เวลาที่สร้าง"],
    ], [2337, 2337, 2338, 2338])

    # Replace API specification with endpoints implemented by the final backend.
    api_tables = {
        "10.1 Authentication": [
            ["POST", "/api/auth/login", "Public"], ["POST", "/api/auth/logout", "ทุก Role"],
            ["GET", "/api/auth/me", "ทุก Role"], ["POST", "/api/auth/change-password", "ทุก Role"],
            ["POST", "/api/auth/forgot-password/request-otp", "Public"],
            ["POST", "/api/auth/forgot-password/verify-otp", "Public"],
            ["POST", "/api/auth/reset-password", "Public พร้อม Reset Token"],
            ["GET", "/api/profile", "ทุก Role หลังเปลี่ยน Temporary Password"],
            ["PUT", "/api/profile", "ทุก Role เฉพาะ Profile ของตนเอง"],
        ],
        "10.2 User Management": [
            ["GET", "/api/admin/users", "Admin"], ["GET", "/api/admin/users/{userId}", "Admin"],
            ["GET", "/api/admin/employees/available-for-account", "Admin"],
            ["POST", "/api/admin/users", "Admin"], ["PUT", "/api/admin/users/{userId}", "Admin"],
            ["PATCH", "/api/admin/users/{userId}/status", "Admin"],
            ["POST", "/api/admin/users/{userId}/reset-password", "Admin"],
        ],
        "10.3 Employee Management": [
            ["GET", "/api/hr/employees", "HR / Admin"], ["POST", "/api/hr/employees", "HR / Admin"],
            ["GET", "/api/hr/employees/{employeeId}", "HR / Admin"], ["PUT", "/api/hr/employees/{employeeId}", "HR / Admin"],
            ["PATCH", "/api/hr/employees/{employeeId}/status", "HR / Admin"],
        ],
        "10.4 Department Management": [
            ["GET", "/api/hr/departments", "HR / Admin"], ["POST", "/api/hr/departments", "HR / Admin"],
            ["GET", "/api/hr/departments/{departmentId}", "HR / Admin"], ["PUT", "/api/hr/departments/{departmentId}", "HR / Admin"],
            ["PATCH", "/api/hr/departments/{departmentId}/status", "HR / Admin"],
        ],
        "10.5 Position Management": [
            ["GET", "/api/hr/positions", "HR / Admin"], ["POST", "/api/hr/positions", "HR / Admin"],
            ["GET", "/api/hr/positions/{positionId}", "HR / Admin"], ["PUT", "/api/hr/positions/{positionId}", "HR / Admin"],
            ["PATCH", "/api/hr/positions/{positionId}/status", "HR / Admin"],
        ],
        "10.6 Leave Type Management": [
            ["GET", "/api/hr/leave-types", "HR / Admin"], ["POST", "/api/hr/leave-types", "HR / Admin"],
            ["GET", "/api/hr/leave-types/{leaveTypeId}", "HR / Admin"], ["PUT", "/api/hr/leave-types/{leaveTypeId}", "HR / Admin"],
            ["PATCH", "/api/hr/leave-types/{leaveTypeId}/status", "HR / Admin"],
        ],
        "10.7 Leave Entitlement Management": [
            ["GET", "/api/leave/balance", "ทุก Role เฉพาะของตนเอง"], ["GET", "/api/hr/leave-entitlements", "HR / Admin"],
            ["POST", "/api/hr/leave-entitlements", "HR / Admin"], ["GET", "/api/hr/leave-entitlements/{entitlementId}", "HR / Admin"],
            ["PUT", "/api/hr/leave-entitlements/{entitlementId}", "HR / Admin"],
        ],
        "10.8 Leave Request Management": [
            ["GET", "/api/leave/options", "ทุก Role"], ["GET", "/api/leave/requests", "ทุก Role เฉพาะของตนเอง"],
            ["GET", "/api/leave/requests/{requestId}", "Owner"], ["POST", "/api/leave/requests/drafts", "ทุก Role"],
            ["PUT", "/api/leave/requests/{requestId}/draft", "Owner เฉพาะ Draft"],
            ["POST", "/api/leave/requests/submit", "ทุก Role"], ["POST", "/api/leave/requests/{requestId}/submit", "Owner เฉพาะ Draft"],
            ["DELETE", "/api/leave/requests/{requestId}/draft", "Owner เฉพาะ Draft"],
            ["PATCH", "/api/leave/requests/{requestId}/cancel", "Owner เฉพาะ Pending"],
        ],
        "10.9 Approval": [
            ["GET", "/api/supervisor/approvals", "Supervisor"], ["GET", "/api/supervisor/approvals/{requestId}", "Direct Supervisor"],
            ["POST", "/api/supervisor/approvals/{requestId}/decision", "Direct Supervisor"],
        ],
        "10.10 Attachment": [
            ["GET", "/api/leave-attachments/{attachmentId}", "Owner / Direct Supervisor / HR / Admin"],
            ["DELETE", "/api/leave-attachments/{attachmentId}", "Owner เฉพาะ Draft"],
        ],
        "10.11 Holiday Management": [
            ["GET", "/api/hr/holidays", "HR / Admin"], ["POST", "/api/hr/holidays", "HR / Admin"],
            ["GET", "/api/hr/holidays/{holidayId}", "HR / Admin"], ["PUT", "/api/hr/holidays/{holidayId}", "HR / Admin"],
            ["DELETE", "/api/hr/holidays/{holidayId}", "HR / Admin"],
        ],
        "10.13 Reports": [
            ["GET", "/api/supervisor/team-report", "Supervisor เฉพาะ Direct Team"],
            ["GET", "/api/reports/leave-requests", "HR / Admin"],
        ],
        "10.14 Audit Log": [["GET", "/api/admin/audit-logs", "Admin"]],
    }
    for heading, rows in api_tables.items():
        rewrite_table(table_after_heading(document, heading), ["Method", "Endpoint", "Permission"], rows)

    screens = table_after_heading(document, "12. หน้าจอที่ต้องพัฒนา")
    append_table_rows(screens, [
        ["Forgot Password / Verify OTP / Reset Password", "ทุก Role ก่อน Login"],
        ["Profile", "ทุก Role"],
        ["Change Password / Force Change Password", "ทุก Role"],
    ])

    # Current notification workflow has no cancel event.
    replace_exact(document, "INSERT notifications ให้ Supervisor ที่เกี่ยวข้อง", "ไม่สร้าง Notification สำหรับการ Cancel ใน Workflow ปัจจุบัน")
    replace_exact(document, "And แจ้งเตือน Supervisor ที่เกี่ยวข้อง", "And ไม่สร้าง Notification เพิ่มสำหรับการ Cancel")
    replace_exact(document, "Then ระบบต้องแสดงเฉพาะ metadata", "Then ระบบต้องอนุญาตให้เปิดหรือดาวน์โหลดไฟล์ตามสิทธิ์ของ Backend")
    replace_exact(document, "And ไม่อนุญาตให้ download หรือเปิดไฟล์", "And ต้องตรวจ Authentication, Role และ Data Scope ก่อนส่งไฟล์")

    # Audit log wording must match the implemented event list, not claim every action.
    replace_exact(document, "ทุกการเปลี่ยนสถานะสำคัญต้องบันทึก Audit Log", "ระบบบันทึก Event สำคัญที่กำหนดใน Audit Log ตามจุดที่ Backend รองรับ")
    replace_exact(document, "ระบบต้องบันทึกกิจกรรมสำคัญ", "ระบบต้องบันทึก Event สำคัญที่กำหนด และไม่อ้างว่าครอบคลุมทุก Action")
    replace_paragraph_range(document, "เหตุการณ์ที่ต้องบันทึก", "ข้อมูลที่ต้องบันทึก", [
        "user_created",
        "admin_password_reset",
        "change_password",
        "password_reset_rate_limited",
        "password_reset_otp_requested",
        "password_reset_otp_failed",
        "password_reset_otp_verified",
        "password_reset_completed",
        "leave_approved",
        "leave_rejected",
        "หมายเหตุ: Backend ปัจจุบันบันทึกเฉพาะ Event สำคัญข้างต้น ไม่อ้างว่าบันทึกทุก Action ของระบบ",
    ])

    # Add explicit Future Scope as the final section without removing Acceptance Criteria.
    anchor = document.paragraphs[-1]
    normal_template = find_paragraph(document, "15. Acceptance Criteria หลัก")
    bullet_template = find_paragraph(document, "ผู้ใช้ Logout ได้")
    heading = deepcopy(normal_template._p)
    document._element.body.insert(document._element.body.index(document._element.body.sectPr), heading)
    new_heading = document.paragraphs[-1]
    set_text(new_heading, "16. Future Scope")
    for text in [
        "Push Notification และ Real-time Notification", "Native Mobile Application", "Multi-step Approval",
        "HR Approval หรือ Final Approval", "Department Manager Approval", "Backup หรือ Delegate Approver",
        "การยกเลิกหลัง Approved พร้อมคืนสิทธิ์วันลา", "Half-day Leave", "Cross-year Leave Request",
        "PDF Export และ Advanced Analytics", "Backup / Restore และ Large-scale Load Test",
    ]:
        node = deepcopy(bullet_template._p)
        document._element.body.insert(document._element.body.index(document._element.body.sectPr), node)
        paragraph = document.paragraphs[-1]
        set_text(paragraph, text)

    # Clean any comparison/history language inherited from the master.
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "เหตุผลที่ไม่ทำในระบบเวอร์ชันปัจจุบัน":
            set_text(paragraph, "สถานะ")
        text = paragraph.text
        if "MVP" in text:
            set_text(paragraph, text.replace("MVP", "ระบบเวอร์ชันปัจจุบัน"))
        if "ใน ระบบเวอร์ชันปัจจุบัน" in paragraph.text:
            set_text(paragraph, paragraph.text.replace("ใน ระบบเวอร์ชันปัจจุบัน", "ในระบบเวอร์ชันปัจจุบัน"))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "MVP" in paragraph.text:
                        set_text(paragraph, paragraph.text.replace("MVP", "ระบบเวอร์ชันปัจจุบัน"))
                    if "ใน ระบบเวอร์ชันปัจจุบัน" in paragraph.text:
                        set_text(paragraph, paragraph.text.replace("ใน ระบบเวอร์ชันปัจจุบัน", "ในระบบเวอร์ชันปัจจุบัน"))

    # Preserve the existing table appearance while identifying all first rows as
    # repeating header rows for multipage readability and accessibility.
    for table in document.tables:
        if not table.rows:
            continue
        grid = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst]
        if grid:
            total = sum(grid)
            tbl_pr = table._tbl.tblPr
            for existing in list(tbl_pr.findall(qn("w:tblW"))):
                tbl_pr.remove(existing)
            tbl_w = OxmlElement("w:tblW")
            tbl_w.set(qn("w:w"), str(total))
            tbl_w.set(qn("w:type"), "dxa")
            tbl_pr.insert(0, tbl_w)
            for existing in list(tbl_pr.findall(qn("w:tblInd"))):
                tbl_pr.remove(existing)
            tbl_ind = OxmlElement("w:tblInd")
            tbl_ind.set(qn("w:w"), "120")
            tbl_ind.set(qn("w:type"), "dxa")
            tbl_pr.insert(1, tbl_ind)
            for row in table.rows:
                for index, cell in enumerate(row.cells):
                    if index >= len(grid):
                        continue
                    tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        cell._tc.get_or_add_tcPr().append(tc_w)
                    tc_w.set(qn("w:w"), str(grid[index]))
                    tc_w.set(qn("w:type"), "dxa")
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)

    document.core_properties.title = "ระบบอนุมัติใบลาออนไลน์ Final Requirement"
    document.core_properties.subject = "Final Requirement for Online Leave Approval System"
    document.save(WORKING)
    shutil.copy2(WORKING, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
