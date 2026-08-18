from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, document)


def run_format(run):
    rpr = run._element.rPr
    fonts = rpr.rFonts if rpr is not None else None
    return {
        "text": run.text,
        "font": run.font.name,
        "ascii": fonts.get(qn("w:ascii")) if fonts is not None else None,
        "eastAsia": fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        "size_pt": run.font.size.pt if run.font.size else None,
        "bold": run.bold,
        "italic": run.italic,
        "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
    }


def paragraph_data(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": float(fmt.line_spacing) if isinstance(fmt.line_spacing, float) else (
            fmt.line_spacing.pt if fmt.line_spacing else None
        ),
        "left_indent_in": fmt.left_indent.inches if fmt.left_indent else None,
        "first_line_indent_in": fmt.first_line_indent.inches if fmt.first_line_indent else None,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": [run_format(run) for run in paragraph.runs],
    }


def package_inventory(path):
    inventory = []
    with zipfile.ZipFile(path) as archive:
        for info in sorted(archive.infolist(), key=lambda item: item.filename):
            data = archive.read(info.filename)
            inventory.append({
                "path": info.filename,
                "size": len(data),
                "sha256": hashlib.sha256(data).hexdigest(),
            })
    return inventory


def main():
    source = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document(source)
    blocks = []
    table_summaries = []
    font_counter = Counter()
    size_counter = Counter()
    style_counter = Counter()

    for index, (kind, item) in enumerate(iter_block_items(document)):
        if kind == "paragraph":
            data = paragraph_data(item)
            data.update({"index": index, "kind": kind})
            blocks.append(data)
            style_counter[data["style"]] += 1
            for run in item.runs:
                if run.text:
                    font_counter[run.font.name or "<inherited>"] += len(run.text)
                    size_counter[str(run.font.size.pt if run.font.size else "<inherited>")] += len(run.text)
        else:
            rows = []
            for row_index, row in enumerate(item.rows):
                cells = []
                for col_index, cell in enumerate(row.cells):
                    paragraphs = [paragraph_data(p) for p in cell.paragraphs]
                    cells.append({"col": col_index, "text": "\n".join(p.text for p in cell.paragraphs), "paragraphs": paragraphs})
                    for p in cell.paragraphs:
                        style_counter[p.style.name if p.style else None] += 1
                        for run in p.runs:
                            if run.text:
                                font_counter[run.font.name or "<inherited>"] += len(run.text)
                                size_counter[str(run.font.size.pt if run.font.size else "<inherited>")] += len(run.text)
                rows.append({"row": row_index, "cells": cells})
            tbl_grid = item._tbl.tblGrid
            grid = [int(col.get(qn("w:w"))) for col in tbl_grid.gridCol_lst] if tbl_grid is not None else []
            table_data = {
                "index": index,
                "kind": kind,
                "table_index": len(table_summaries),
                "style": item.style.name if item.style else None,
                "rows": rows,
                "grid_dxa": grid,
            }
            blocks.append(table_data)
            table_summaries.append({
                "table_index": len(table_summaries),
                "block_index": index,
                "style": table_data["style"],
                "row_count": len(rows),
                "column_count": max((len(r["cells"]) for r in rows), default=0),
                "grid_dxa": grid,
                "first_row": [cell["text"] for cell in rows[0]["cells"]] if rows else [],
            })

    metadata = {
        "source": str(source),
        "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "block_count": len(blocks),
        "styles": style_counter.most_common(),
        "fonts_by_characters": font_counter.most_common(),
        "sizes_by_characters": size_counter.most_common(),
        "tables": table_summaries,
    }
    (output_dir / "reference-metadata.json").write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "reference-blocks.json").write_text(json.dumps(blocks, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "reference-package-inventory.json").write_text(json.dumps(package_inventory(source), ensure_ascii=False, indent=2), encoding="utf-8")

    lines = []
    for block in blocks:
        if block["kind"] == "paragraph":
            text = block["text"].replace("\n", " ").strip()
            if text:
                lines.append(f"P{block['index']:04d} [{block['style']}] {text}")
        else:
            lines.append(f"T{block['table_index']:03d} BLOCK={block['index']} ROWS={len(block['rows'])} GRID={block['grid_dxa']}")
            for row in block["rows"]:
                values = [cell["text"].replace("\n", " / ").strip() for cell in row["cells"]]
                lines.append(f"  R{row['row']:03d} | " + " | ".join(values))
    (output_dir / "reference-content.txt").write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
