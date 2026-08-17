from __future__ import annotations

import hashlib
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT = "TH Sarabun New"
BODY_SIZE = 16
TABLE_SIZE = 12
CONTENT_WIDTH_DXA = 9350


def set_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_para(document, text="", *, bold=False, size=BODY_SIZE, align=None,
             before=0, after=6, keep=False, page_break=False, italic=False):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = keep
    paragraph.paragraph_format.page_break_before = page_break
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph.add_run(text), size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(document, text, level=1, page_break=False):
    if level == 1:
        return add_para(document, text, bold=True, size=18, before=10, after=8,
                        keep=True, page_break=page_break)
    return add_para(document, text, bold=True, size=16, before=8, after=5, keep=True)


def add_bullet(document, text, level=0):
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.style = document.styles["List Paragraph"]
    paragraph.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.10
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    set_font(paragraph.add_run(text), size=BODY_SIZE)
    return paragraph


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths, center_columns=()):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for col, value in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(str(value)), size=TABLE_SIZE, bold=True)
        set_cell_margins(cell)
    for values in rows:
        cells = table.add_row().cells
        for col, value in enumerate(values):
            cell = cells[col]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_font(p.add_run(str(value)), size=TABLE_SIZE)
            set_cell_margins(cell)
    set_table_geometry(table, widths)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def build(source: Path, output: Path):
    expected = "dc7b176ce7cb5462ef26f5d8e1faa60444f742b0be9541c82ddfe90f671fbc0a"
    actual = hashlib.sha256(source.read_bytes()).hexdigest()
    if actual != expected:
        raise RuntimeError("Reference DOCX hash changed; fresh distillation is required.")
    shutil.copy2(source, output)
    document = Document(output)
    clear_body(document)
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_para(document, "ระบบอนุมัติใบลาออนไลน์", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=3)
    add_para(document, "Online Leave Approval System", size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_para(document, "Final Requirement", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    add_heading(document, "1. สรุปผลการเปรียบเทียบ Requirement เดิมกับระบบ Final")
    add_para(document, "เอกสารเดิมสอดคล้องกับระบบในสาระหลัก ได้แก่ ผู้ใช้งาน 4 บทบาท การยื่นคำขอลาของทุกบทบาท การอนุมัติโดย Supervisor ขั้นเดียว และ Function ด้าน HR/Admin อย่างไรก็ตาม ต้องปรับรายละเอียดด้าน Authentication, Profile, Supervisor Data Scope, Attachment, Notification และ Audit Log ให้ตรงกับระบบ Final ที่ผ่านการทดสอบแล้ว")
    add_para(document, "การปรับครั้งนี้คงรหัส FR-01 ถึง FR-16 เท่าที่เหมาะสม และเพิ่ม FR-17 ถึง FR-23 สำหรับความสามารถที่เอกสารเดิมยังระบุไม่ครบ ส่วน Multi-step Approval, HR/Department Manager Approval, Backup Approver, Push/Real-time Notification และ Native Mobile Application ถูกแยกไว้ใน Future Scope")
    add_para(document, "ผลการจัดประเภทสำหรับ Review", bold=True, keep=True)
    add_table(document, ["สถานะ", "ความหมาย", "ผลการทบทวน"], [
        ("KEEP", "สาระเดิมตรงกับระบบ Final", "คงไว้และปรับถ้อยคำเพื่อความชัดเจน"),
        ("UPDATE", "แนวคิดเดิมถูกต้องแต่รายละเอียดไม่ตรงทั้งหมด", "แก้ Role, Data Scope, Validation หรือพฤติกรรมให้ตรงระบบ"),
        ("ADD", "ระบบ Final มีความสามารถจริงแต่เอกสารเดิมยังไม่ระบุ", "เพิ่ม Requirement ใหม่ในหมวดที่เกี่ยวข้อง"),
        ("REMOVE", "ข้อความเดิมไม่ใช่ Current Requirement", "ตัดออกจาก Current Scope พร้อมระบุเหตุผล"),
        ("FUTURE SCOPE", "ยังไม่มีในระบบ Final", "แยกจาก Current Requirement"),
        ("GAP", "เอกสารเดิมกำหนดไว้แต่ระบบ Final ยังไม่มีพฤติกรรมนั้น", "บันทึกเพื่อให้เจ้าของระบบยืนยัน ไม่แก้ระบบในรอบนี้"),
    ], [1500, 3350, 4500], center_columns=(0,))

    add_heading(document, "2. ขอบเขตและบทบาทของระบบ")
    add_para(document, "Online Leave Approval System เป็นระบบสำหรับบริหารจัดการกระบวนการลาของบุคลากรภายในองค์กร ตั้งแต่การสร้างและส่งคำขอลา การตรวจสอบสิทธิ์และวันลาคงเหลือ การคำนวณวันทำงาน การอนุมัติหรือปฏิเสธโดยหัวหน้างาน การติดตามสถานะ การแจ้งเตือนภายในระบบ ตลอดจนการจัดการข้อมูลพนักงาน ประเภทการลา สิทธิ์การลา วันหยุด บัญชีผู้ใช้ หน่วยงาน ตำแหน่ง รายงาน และประวัติการใช้งาน")
    add_para(document, "ระบบมีผู้ใช้งาน 4 บทบาท ได้แก่ Employee, Supervisor, HR และ Admin โดยผู้ใช้งานทุกบทบาทยังสามารถใช้ความสามารถด้านคำขอลาของตนเองได้ตามเงื่อนไขของข้อมูลและสิทธิ์การลา ส่วน Workflow การอนุมัติปัจจุบันมีเพียงขั้นเดียว คือ ผู้ใช้งานส่งคำขอลาให้ Supervisor ซึ่งเป็นหัวหน้างานโดยตรงพิจารณาอนุมัติหรือปฏิเสธ")
    add_table(document, ["Role", "หน้าที่หลัก"], [
        ("Employee", "ยื่นและติดตามคำขอลา ตรวจสอบสิทธิ์การลา จัดการ Notification และ Profile ของตนเอง"),
        ("Supervisor", "ใช้ความสามารถด้านการลาของตนเอง และพิจารณาคำขอของพนักงานที่อยู่ภายใต้การดูแลโดยตรง รวมถึงดู Team Report"),
        ("HR", "ใช้ความสามารถด้านการลาของตนเอง และจัดการ Employee, Leave Type, Leave Entitlement, Holiday และ HR Report"),
        ("Admin", "ใช้ความสามารถด้านการลาของตนเอง และจัดการ User, Department, Position และ Audit Log"),
    ], [2200, 7150], center_columns=(0,))

    add_heading(document, "3. Functional Requirement Final", page_break=True)
    functional = [
        ("FR-01", "ระบบสามารถให้ผู้ใช้งานเข้าสู่ระบบด้วย Username หรือ Email และ Password ออกจากระบบ ตรวจสอบ Session และควบคุมการเข้าถึงตาม Role ได้ รวมถึงบังคับเปลี่ยน Temporary Password ก่อนใช้งาน Function ธุรกิจ และให้ผู้ใช้เปลี่ยน Password ของตนเองตาม Password Policy", "ทุก Role"),
        ("FR-02", "ระบบสามารถให้ Admin เพิ่ม แก้ไข Role, Username และสถานะบัญชีผู้ใช้ สร้าง Temporary Password และ Reset Password ของบัญชีที่อนุญาตได้ โดยบัญชีหนึ่งผูกกับ Employee ได้เพียงหนึ่งราย", "Admin"),
        ("FR-03", "ระบบสามารถให้ HR และ Admin ดู เพิ่ม แก้ไข ค้นหา และเปลี่ยนสถานะ Employee รวมถึงกำหนด Department, Position และ Supervisor ได้", "HR, Admin"),
        ("FR-04", "ระบบสามารถให้ Admin จัดการข้อมูล Department โดยเพิ่ม แก้ไข และเปลี่ยนสถานะ Active หรือ Inactive ได้", "Admin"),
        ("FR-05", "ระบบสามารถให้ Admin จัดการข้อมูล Position โดยเพิ่ม แก้ไข และเปลี่ยนสถานะ Active หรือ Inactive ได้", "Admin"),
        ("FR-06", "ระบบสามารถให้ HR และ Admin ดู เพิ่ม แก้ไข และเปลี่ยนสถานะ Leave Type พร้อมกำหนด Code, Quota, Minimum/Maximum Days และ Attachment Policy ได้", "HR, Admin"),
        ("FR-07", "ระบบสามารถให้ HR และ Admin ดู เพิ่ม และแก้ไข Leave Entitlement ของ Employee แยกตาม Leave Type และปีได้", "HR, Admin"),
        ("FR-08", "ระบบสามารถให้ผู้ใช้งานทุก Role สร้าง Leave Request บันทึกเป็น Draft แก้ไขหรือลบ Draft และ Submit คำขอของตนเองได้ โดยระบบสร้าง Request Number เมื่อ Submit สำเร็จ", "ทุก Role"),
        ("FR-09", "ระบบสามารถตรวจสอบข้อมูล Leave Request ก่อน Submit ได้แก่ Required Field, Date Range, Same Year, Reason Length, Active Leave Type, Working Days, Minimum/Maximum Days, Entitlement, Balance, Date Overlap, Attachment Policy และการมี Supervisor Account", "ทุก Role"),
        ("FR-10", "ระบบสามารถให้ Supervisor ดูรายละเอียดและ Approve หรือ Reject Leave Request สถานะ Pending ของพนักงานที่อยู่ภายใต้การดูแลโดยตรงได้ โดยห้ามพิจารณาคำขอของตนเอง และ Reject ต้องระบุเหตุผล", "Supervisor"),
        ("FR-11", "ระบบสามารถให้เจ้าของยกเลิก Leave Request ของตนเองที่อยู่ในสถานะ Pending ได้ โดยเปลี่ยนสถานะเป็น Cancelled", "ทุก Role"),
        ("FR-12", "ระบบสามารถรับและจัดเก็บ Attachment ประกอบ Leave Request ตาม Policy ของ Leave Type พร้อมควบคุมชนิด จำนวน ขนาด และสิทธิ์การดาวน์โหลดหรือลบไฟล์", "ทุก Role ตาม Data Scope"),
        ("FR-13", "ระบบสามารถให้ HR และ Admin ดู เพิ่ม แก้ไข และลบ Holiday ได้ และนำเฉพาะ Holiday ที่ Active ไปใช้ในการคำนวณ Working Days", "HR, Admin"),
        ("FR-14", "ระบบสามารถให้ Supervisor ดู Team Report ของ Direct Team และให้ HR หรือ Admin ดู HR Leave Report พร้อม Filter และ Summary ตามข้อมูลที่ระบบรองรับ โดย HR Report สามารถส่งออกเป็น Excel ได้", "Supervisor, HR, Admin"),
        ("FR-15", "ระบบสามารถสร้าง In-app Notification เมื่อมีการ Submit เพื่อแจ้ง Supervisor และเมื่อมีการ Approve หรือ Reject เพื่อแจ้งเจ้าของคำขอได้", "ทุก Role ตาม Event และ Ownership"),
        ("FR-16", "ระบบสามารถบันทึก Event สำคัญที่ระบบกำหนดและให้ Admin ดู Audit Log ได้ โดยการเข้าถึง Audit Log จำกัดเฉพาะ Admin", "Admin, System"),
        ("FR-17", "ระบบสามารถให้ผู้ใช้งานขอ OTP ทาง Email ยืนยัน OTP และ Reset Password ด้วย Reset Token แบบใช้ครั้งเดียวภายใต้ข้อจำกัดด้านเวลาและจำนวนครั้งได้", "ทุก Role"),
        ("FR-18", "ระบบสามารถให้ผู้ใช้งานดูและแก้ไข Profile ของตนเอง รวมถึงอัปโหลด เปลี่ยน หรือลบ Profile Image ตามเงื่อนไขของระบบได้", "ทุก Role"),
        ("FR-19", "ระบบสามารถให้ผู้ใช้งานดูรายการและรายละเอียด Leave Request ของตนเองทุกสถานะ พร้อมข้อมูลการลาและ Attachment ที่เกี่ยวข้องได้", "ทุก Role"),
        ("FR-20", "ระบบสามารถแสดง Leave Balance ของผู้ใช้งานแยกตาม Leave Type และปี โดยแสดง Total, Used, Pending และ Remaining ได้", "ทุก Role"),
        ("FR-21", "ระบบสามารถให้ Supervisor ดูรายการ Pending Approval และ Team Report เฉพาะพนักงานที่อยู่ภายใต้การดูแลโดยตรง พร้อม Filter และ Summary ที่ระบบรองรับได้", "Supervisor"),
        ("FR-22", "ระบบสามารถให้ผู้ใช้งานดูจำนวน Unread, Mark Read, Mark All Read, Delete Notification ของตนเอง และเปิด Deep Link ไปยังรายการที่เกี่ยวข้องได้", "ทุก Role"),
        ("FR-23", "ระบบสามารถแสดง Dashboard ตาม Role โดยใช้ข้อมูลจากระบบกลางและเชื่อมโยงไปยัง Function ที่ Role นั้นมีสิทธิ์ใช้งานได้", "ทุก Role"),
    ]
    add_table(document, ["รหัส", "Functional Requirement", "Role ที่เกี่ยวข้อง"], functional,
              [1050, 6450, 1850], center_columns=(0, 2))

    add_heading(document, "4. Role and Permission Requirement", page_break=True)
    add_para(document, "การควบคุมสิทธิ์ต้องตรวจทั้งการเข้าถึงหน้าจอและ Backend API โดย Backend เป็นจุดบังคับสิทธิ์สุดท้าย ข้อมูลส่วนบุคคล คำขอลา และ Notification ต้องจำกัดตามเจ้าของข้อมูล ส่วน Supervisor ต้องถูกจำกัดเฉพาะ Direct Team และไม่สามารถพิจารณาคำขอของตนเองได้")
    matrix = [
        ("Login, Logout, Profile และ Change Password", "✓", "✓", "✓", "✓"),
        ("Own Leave: Draft, Submit, List, Detail และ Cancel", "เงื่อนไข", "เงื่อนไข", "เงื่อนไข", "เงื่อนไข"),
        ("ดู Leave Balance ของตนเอง", "✓", "✓", "✓", "✓"),
        ("จัดการ Notification ของตนเอง", "✓", "✓", "✓", "✓"),
        ("Pending Approval และ Decision", "–", "เงื่อนไข", "–", "–"),
        ("Team Report", "–", "เงื่อนไข", "–", "–"),
        ("Employee Management", "–", "–", "✓", "✓"),
        ("Leave Type / Entitlement / Holiday", "–", "–", "✓", "✓"),
        ("HR Leave Report", "–", "–", "✓", "✓"),
        ("User Management", "–", "–", "–", "✓"),
        ("Department / Position Management UI", "–", "–", "–", "✓"),
        ("Department / Position API", "–", "–", "✓", "✓"),
        ("ดู Audit Log", "–", "–", "–", "✓"),
        ("ดาวน์โหลด Attachment ของตนเอง", "✓", "✓", "✓", "✓"),
        ("ดาวน์โหลด Attachment ของ Direct Team", "–", "เงื่อนไข", "–", "–"),
        ("ดาวน์โหลด Attachment ตามสิทธิ์ HR/Admin", "–", "–", "✓", "✓"),
    ]
    add_table(document, ["Function", "Employee", "Supervisor", "HR", "Admin"], matrix,
              [3750, 1400, 1400, 1400, 1400], center_columns=(1, 2, 3, 4))
    add_para(document, "คำอธิบายเงื่อนไข", bold=True, keep=True)
    add_bullet(document, "Own Leave Submit ต้องมี Employee Record, Leave Entitlement ของปีและประเภทลา, Balance เพียงพอ และ Supervisor ที่มีบัญชีผู้ใช้")
    add_bullet(document, "Supervisor พิจารณาและดูรายงานได้เฉพาะ Direct Team และไม่รวมคำขอของตนเอง")
    add_bullet(document, "Notification และข้อมูล Own Leave จำกัดด้วย User หรือ Employee ของ Session")
    add_bullet(document, "Attachment ดาวน์โหลดได้เฉพาะเจ้าของ Direct Supervisor, HR หรือ Admin ส่วนการลบจำกัดเฉพาะเจ้าของ Draft")

    add_heading(document, "5. Business Rule Final")
    rules = [
        ("BR-01", "การ Submit ต้องระบุ Leave Type, Start Date, End Date และ Reason ครบถ้วน"),
        ("BR-02", "End Date ต้องไม่ก่อน Start Date และวันที่ต้องอยู่ในรูปแบบที่ระบบยอมรับ"),
        ("BR-03", "Leave Request หนึ่งรายการต้องไม่ข้ามปี"),
        ("BR-04", "Reason ต้องมีความยาวตั้งแต่ 5 ถึง 500 ตัวอักษร"),
        ("BR-05", "Submit ได้เฉพาะ Leave Type ที่ Active"),
        ("BR-06", "Working Days ไม่นับวันเสาร์ วันอาทิตย์ และ Holiday ที่ Active"),
        ("BR-07", "ช่วงวันที่ที่ไม่มี Working Day ไม่สามารถ Submit ได้"),
        ("BR-08", "จำนวนวันลาต้องอยู่ระหว่าง Minimum Days และ Maximum Days Per Request ของ Leave Type"),
        ("BR-09", "ผู้ยื่นต้องมี Leave Entitlement ของ Employee, Leave Type และปีที่เลือก"),
        ("BR-10", "Available Balance สำหรับการยื่นคำนวณจาก Total ลบ Used และ Pending"),
        ("BR-11", "Pending Request ของประเภทและปีเดียวกันต้องถูกนำไปหักออกจาก Available Balance"),
        ("BR-12", "ห้ามช่วงวันลาทับกับ Leave Request สถานะ Pending หรือ Approved ของ Employee คนเดียวกัน"),
        ("BR-13", "เมื่อ Leave Type กำหนดให้แนบเอกสารและถึง Threshold ผู้ยื่นต้องมี Attachment ก่อน Submit"),
        ("BR-14", "Attachment รองรับ PDF, JPEG และ PNG สูงสุด 5 ไฟล์ โดยแต่ละไฟล์ไม่เกิน 10 MB"),
        ("BR-15", "ดาวน์โหลด Attachment ได้เฉพาะเจ้าของ Direct Supervisor, HR หรือ Admin"),
        ("BR-16", "แก้ไขหรือลบ Leave Request และลบ Attachment ได้เฉพาะรายการที่เป็น Owned Draft"),
        ("BR-17", "เจ้าของยกเลิกได้เฉพาะ Leave Request สถานะ Pending"),
        ("BR-18", "Submit ต้องพบ Supervisor ที่ผูกกับ Employee และ Supervisor ต้องมี User Account"),
        ("BR-19", "เฉพาะ Role Supervisor เท่านั้นที่ใช้ Approval Function ได้ และคำขอต้องเป็นของ Direct Team"),
        ("BR-20", "Supervisor ห้าม Approve หรือ Reject Leave Request ของตนเอง"),
        ("BR-21", "Approve หรือ Reject ทำได้เฉพาะสถานะ Pending และคำขอที่ตัดสินแล้วไม่สามารถตัดสินซ้ำ"),
        ("BR-22", "Reject ต้องระบุ Rejection Reason"),
        ("BR-23", "Approve ต้องตรวจ Entitlement และ Balance ซ้ำภายใน Transaction ก่อนเพิ่ม Used Days"),
        ("BR-24", "Employee Code และ Email ต้องไม่ซ้ำ และข้อมูลอ้างอิง Department, Position และ Supervisor ต้องถูกต้อง"),
        ("BR-25", "Department Name และ Position Name ต้องมีความยาว 2 ถึง 100 ตัวอักษรและไม่ซ้ำโดยไม่คำนึงถึงตัวพิมพ์"),
        ("BR-26", "Leave Type Code ต้องเป็นตัวพิมพ์ใหญ่หรือตัวเลข 2 ถึง 10 ตัว และ Code กับ Name ต้องไม่ซ้ำ"),
        ("BR-27", "Holiday Date ต้องไม่ซ้ำ และข้อมูล Name, Type, Description และ Status ต้องผ่านการตรวจสอบ"),
        ("BR-28", "Leave Entitlement ต้องไม่ซ้ำในชุด Employee, Leave Type และ Year โดย Year อยู่ระหว่าง 2000 ถึง 2100 และ Total/Used Days อยู่ระหว่าง 0 ถึง 365"),
        ("BR-29", "Username ต้องไม่ซ้ำ มีความยาว 4 ถึง 50 ตัว และใช้ตัวพิมพ์เล็ก ตัวเลข จุด ขีดล่าง หรือขีดกลางตามรูปแบบที่กำหนด"),
        ("BR-30", "User Account หนึ่งรายการผูกกับ Employee หนึ่งรายที่ยังไม่มีบัญชี Role ต้อง Active และสถานะบัญชีต้องเป็น Active, Inactive หรือ Locked"),
        ("BR-31", "Admin Create หรือ Reset Password ต้องสร้าง Temporary Password และบังคับให้ผู้ใช้เปลี่ยน Password โดยบัญชี Inactive ไม่สามารถ Reset Password ได้"),
        ("BR-32", "Password ใหม่ต้องมีอย่างน้อย 10 ตัวอักษร มีตัวพิมพ์ใหญ่ ตัวพิมพ์เล็ก ตัวเลข อักขระพิเศษ ไม่มีช่องว่าง และไม่เท่ากับ Username, Email หรือ Password เดิม"),
        ("BR-33", "Profile Image รองรับ JPEG, PNG หรือ WebP หนึ่งไฟล์ ขนาดไม่เกิน 2 MB และเนื้อหาไฟล์ต้องตรงกับชนิดไฟล์"),
        ("BR-34", "Notification List, Read และ Delete ต้องจำกัดด้วย User ของ Session"),
        ("BR-35", "Protected Function ต้องปฏิเสธบัญชีที่ไม่ Active และ Session ที่หมดอายุหรือ Token Version ไม่ตรง"),
        ("BR-36", "OTP เป็นตัวเลข 6 หลัก มีอายุเริ่มต้น 5 นาที ตรวจผิดได้สูงสุด 5 ครั้ง มีระยะรอขอใหม่และจำกัดจำนวนการขอในช่วงเวลา"),
        ("BR-37", "Reset Token มีอายุเริ่มต้น 15 นาที ใช้ได้ครั้งเดียว และเมื่อ Reset สำเร็จต้องยกเลิก Token และ OTP อื่นที่ยังใช้ได้ของผู้ใช้รายนั้น"),
    ]
    add_table(document, ["รหัส", "Business Rule"], rules, [1200, 8150], center_columns=(0,))

    add_heading(document, "6. Non-Functional Requirement Final", page_break=True)
    nfr = [
        ("NFR-01", "ระบบต้องรักษาความปลอดภัยของ Session โดยตรวจลายเซ็น อายุ สถานะบัญชี และ Token Version"),
        ("NFR-02", "Browser Session ต้องใช้ HTTP-only Cookie และกำหนด SameSite รวมถึง Secure ตามสภาพแวดล้อม Production"),
        ("NFR-03", "Password และ Temporary Password ต้องจัดเก็บเป็น Hash ส่วน OTP และ Reset Token ต้องไม่จัดเก็บเป็น Plain Text"),
        ("NFR-04", "Backend ต้องบังคับ Role และ Data Scope ของทุก Endpoint สำคัญ แม้มีการเรียก Direct API"),
        ("NFR-05", "ข้อมูลธุรกิจหลักต้องใช้ Backend API และฐานข้อมูลกลางเป็น Source of Truth เดียวกัน"),
        ("NFR-06", "กระบวนการ Submit, Approve, Reject และ Password Reset ที่เปลี่ยนหลายข้อมูลต้องทำงานแบบ Transaction"),
        ("NFR-07", "Backend ต้องตรวจ Required Field, รูปแบบข้อมูล, ช่วงค่า, Duplicate และ Reference ก่อนบันทึก"),
        ("NFR-08", "การ Upload ต้องจำกัดจำนวน ขนาด และชนิดไฟล์ พร้อมตรวจสิทธิ์เข้าถึงไฟล์"),
        ("NFR-09", "ระบบต้องแสดง Error Message ที่เข้าใจได้และไม่เปิดเผยข้อมูลบัญชีหรือรายละเอียดภายในที่ไม่ควรเปิดเผย"),
        ("NFR-10", "ระบบต้องรองรับ Responsive Web และ Function หลักต้องใช้งานได้บน Desktop และ Mobile ขนาด 390 x 844"),
        ("NFR-11", "Interactive Control สำคัญ โดยเฉพาะ Icon-only Action ต้องมี Accessible Name และรองรับ Focus"),
        ("NFR-12", "Event ด้าน Security และ Approval ที่ระบบกำหนดต้องมีข้อมูลผู้กระทำ Action, Record, Result และ Metadata เพื่อการตรวจสอบย้อนหลัง"),
    ]
    add_table(document, ["รหัส", "Non-Functional Requirement"], nfr, [1200, 8150], center_columns=(0,))

    add_heading(document, "7. Future Scope", page_break=True)
    add_para(document, "รายการต่อไปนี้ไม่อยู่ใน Current Requirement และยังไม่มีใน Workflow ปัจจุบัน")
    for item in [
        "Push Notification เข้าโทรศัพท์",
        "Native Mobile Application",
        "Real-time Notification",
        "Multi-step Approval",
        "HR Approval หรือ Final Approval",
        "Department Manager Approval",
        "Backup Approver หรือ Delegate Approval",
        "การยกเลิกหลัง Approved พร้อมกระบวนการคืนสิทธิ์ลา",
        "Half-day Leave และ Cross-year Leave Request",
    ]:
        add_bullet(document, item)
    add_para(document, "Email Service ที่ใช้อยู่ใน Current Scope มีหน้าที่ส่ง OTP สำหรับ Forgot Password เท่านั้น ไม่ใช่ Email Notification สำหรับเหตุการณ์การลา")

    add_heading(document, "8. Change Summary", page_break=True)
    changes = [
        ("FR-01 Authentication", "UPDATE", "เพิ่ม Session, Temporary Password, Password Policy และแยก Forgot Password เป็น FR-17", "เอกสารเดิมยังไม่ครอบคลุม Authentication Flow ปัจจุบัน"),
        ("FR-02 User Management", "UPDATE", "เพิ่ม Temporary Password, Reset Password และกฎบัญชี Inactive", "ตรงกับ Admin User Function ปัจจุบัน"),
        ("FR-03 Employee Management", "UPDATE", "เพิ่ม Admin ใน Backend Permission และปรับกฎข้อมูลซ้ำ/Reference", "HR และ Admin ใช้ HR Management API ได้"),
        ("FR-04 Department Management", "KEEP", "คง Admin Management และเพิ่มกฎชื่อซ้ำ", "Function หลักตรงระบบ Final"),
        ("FR-05 Position Management", "KEEP", "คง Admin Management และเพิ่มกฎชื่อซ้ำ", "Function หลักตรงระบบ Final"),
        ("FR-06 Leave Type Management", "UPDATE", "เพิ่ม HR/Admin, Code, Minimum/Maximum Days และ Policy Validation", "ระบบ Final มี Field และ Validation เพิ่มจากเอกสารเดิม"),
        ("FR-07 Leave Entitlement", "UPDATE", "ตัดข้อความว่าทุกการปรับต้องมี Audit Log และยืนยันสูตร Balance", "Current Audit Coverage ยังไม่ครอบคลุมการแก้ Entitlement ทุก Action"),
        ("FR-08 Leave Request Creation", "UPDATE", "แก้จากการบันทึก approver ตั้งแต่ Submit เป็นการตรวจ Direct Supervisor Account", "Current Workflow ตรวจ Direct Team จากความสัมพันธ์ Employee-Supervisor"),
        ("FR-09 Leave Validation", "UPDATE", "เพิ่ม Reason Length, Minimum/Maximum Days และ Attachment Threshold; ตัดกฎลาป่วย 3 วันแบบตายตัว", "Attachment Rule มาจาก Leave Type Policy"),
        ("FR-10 Leave Approval", "UPDATE", "ใช้ Direct Team Scope และ Audit Log; ตัด Approval Log แยก", "Current Decision Flow ไม่เขียนตาราง Approval History แยก"),
        ("FR-11 Leave Cancellation", "KEEP", "คง Owner Cancel เฉพาะ Pending", "ตรงกับ Current API"),
        ("FR-12 Attachment", "UPDATE", "Admin และ HR ดาวน์โหลดได้ตาม Role; จำกัดสูงสุด 5 ไฟล์", "Permission เดิมที่ให้ Admin เห็น Metadata อย่างเดียวไม่ตรงระบบ Final"),
        ("FR-13 Holiday", "UPDATE", "รองรับดู เพิ่ม แก้ไข และลบ Holiday", "Current API ใช้ Delete ไม่ใช่ Status Update Endpoint"),
        ("FR-14 Reporting", "UPDATE", "แยก Team Report และ HR Report พร้อม Admin Access", "Permission และรายงานปัจจุบันต่างจากตารางเดิมบางส่วน"),
        ("FR-15 Notification", "UPDATE", "Event ปัจจุบันมี Submit, Approve, Reject และเพิ่ม Unread/Delete/Deep Link", "ไม่มี Notification Event ตอน Cancel ใน Current Implementation"),
        ("FR-16 Audit Log", "UPDATE", "กำหนดตาม Event ที่ระบบบันทึกจริง ไม่ระบุว่าทุกกิจกรรมถูกบันทึก", "Audit Coverage ปัจจุบันเป็นแบบเลือก Event"),
        ("Forgot Password / OTP", "ADD", "เพิ่ม FR-17", "ระบบ Final มี Email OTP และ Reset Token"),
        ("Profile / Profile Image", "ADD", "เพิ่ม FR-18", "เอกสารเดิมยังไม่มี Requirement นี้"),
        ("My Requests / Detail", "ADD", "เพิ่ม FR-19", "ระบบ Final มี Owner-scoped List และ Detail"),
        ("Leave Balance", "ADD", "เพิ่ม FR-20", "แยกความสามารถแสดง Balance ให้ชัดเจน"),
        ("Supervisor Pending / Team Report", "ADD", "เพิ่ม FR-21", "ระบุ Direct Team Scope และ Summary ให้ชัดเจน"),
        ("Notification Management", "ADD", "เพิ่ม FR-22", "เอกสารเดิมขาด Delete และ Deep Link"),
        ("Role Dashboard", "ADD", "เพิ่ม FR-23", "ระบบ Final มี Dashboard แยกตาม Role"),
        ("Approval Log แยก", "GAP", "ไม่นำมาเป็น Current Requirement", "ระบบ Final ใช้ข้อมูล Approver/Timestamp ใน Leave Request ร่วมกับ Audit Log"),
        ("Cancel Notification/Audit", "GAP", "ไม่นำมาเป็น Current Requirement", "Cancel Endpoint ปัจจุบันเปลี่ยนสถานะเท่านั้น"),
        ("Comprehensive Audit ทุก Action", "GAP", "ลดขอบเขตให้ตรง Event ที่ระบบเขียนจริง", "ระบบ Final ยังไม่ได้บันทึกทุก Action ที่เอกสารเดิมระบุ"),
        ("HR/Department Manager/Multi-step Approval", "FUTURE SCOPE", "ย้ายออกจาก Current Scope", "Workflow Final มี Direct Supervisor ขั้นเดียว"),
        ("Push/Real-time/Native Mobile", "FUTURE SCOPE", "แยกไว้ใน Future Scope", "ไม่มีในระบบ Final ปัจจุบัน"),
    ]
    add_table(document, ["Requirement เดิม", "สถานะ", "สิ่งที่แก้", "เหตุผล"], changes,
              [2100, 1200, 2850, 3200], center_columns=(1,))

    add_heading(document, "9. ประเด็นที่ต้องยืนยันก่อนจัดทำเอกสาร Final", page_break=True)
    questions = [
        "ยืนยันว่าผู้ใช้งานทุก Role สามารถยื่น Leave Request ของตนเองได้ โดยการ Submit ขึ้นอยู่กับการมี Supervisor Account ใช่หรือไม่",
        "ยืนยันว่า Department และ Position Management มีหน้าจอเฉพาะ Admin แม้ Backend API อนุญาต HR และ Admin ใช่หรือไม่",
        "ยืนยันว่า Notification Event ปัจจุบันต้องมีเฉพาะ Submit, Approve และ Reject และไม่ต้องสร้าง Event เมื่อ Cancel ใช่หรือไม่",
        "ยืนยันว่า Approver และเวลาตัดสินใจใน Leave Request ร่วมกับ Audit Log เพียงพอสำหรับ Workflow ชั้นเดียว และไม่ต้องมี Approval History Table แยกใช่หรือไม่",
        "ยืนยันว่า Audit Log แบบเลือกบันทึกเฉพาะ Event ปัจจุบันเป็นขอบเขตที่ยอมรับได้ หรือควรกำหนดรายการ Event เพิ่มเป็น Requirement Gap สำหรับรอบถัดไป",
        "ยืนยันว่า HR และ Admin สามารถดาวน์โหลด Attachment ตาม Role Permission ปัจจุบันได้ แม้เอกสารเดิมกำหนดข้อจำกัดเข้มกว่านี้ใช่หรือไม่",
        "ยืนยัน Leave Type ที่ต้องบังคับ Attachment และ Threshold ที่จะใช้กับข้อมูล Production เพื่อเตรียม Acceptance Test ในอนาคต",
    ]
    for question in questions:
        add_bullet(document, question)

    add_heading(document, "10. สรุปความสอดคล้องของ Final Requirement")
    for item in [
        "ระบบมี Role เพียง Employee, Supervisor, HR และ Admin",
        "Workflow การอนุมัติปัจจุบันเป็น Direct Supervisor ขั้นเดียว",
        "ไม่มี HR Approval, Department Manager Approval, Final Approver หรือ Backup Approver ใน Current Requirement",
        "Push Notification, Real-time Notification และ Native Mobile Application อยู่ใน Future Scope",
        "Functional Requirement, Permission, Business Rule และ Non-Functional Requirement ใช้ข้อมูลและขอบเขตเดียวกับระบบ Final",
        "เอกสารฉบับนี้หยุดที่ Final Requirement และยังไม่รวม Workflow Final, Use Case, DFD, ER Diagram หรือ Data Dictionary",
    ]:
        add_bullet(document, item)

    document.save(output)


if __name__ == "__main__":
    build(Path(sys.argv[1]), Path(sys.argv[2]))
