from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\User\Desktop\online-leavesystem")
DOCX = ROOT / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"
OUT = ROOT / ".codex-docx-audit"
OUT.mkdir(exist_ok=True)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


doc = Document(str(DOCX))
records = []
paragraph_index = 0
table_index = 0

for child in doc.element.body.iterchildren():
    if child.tag == qn("w:p"):
        p = Paragraph(child, doc._body)
        text = norm(p.text)
        records.append({
            "kind": "paragraph",
            "paragraph": paragraph_index,
            "style": p.style.name if p.style else "",
            "text": text,
        })
        paragraph_index += 1
    elif child.tag == qn("w:tbl"):
        table = Table(child, doc._body)
        table_rows = []
        for row_index, row in enumerate(table.rows):
            cells = [norm(cell.text) for cell in row.cells]
            table_rows.append(cells)
            records.append({
                "kind": "table_row",
                "table": table_index,
                "row": row_index,
                "text": " | ".join(cells),
                "cells": cells,
            })
        table_index += 1

texts = [r["text"] for r in records if r["text"]]
exact = defaultdict(list)
for index, record in enumerate(records):
    text = record.get("text", "")
    if len(text) >= 45:
        exact[text.casefold()].append(index)
exact_duplicates = [
    {"text": records[indexes[0]]["text"], "locations": [records[i] for i in indexes]}
    for indexes in exact.values() if len(indexes) > 1
]

id_patterns = {
    "FR": r"\bFR-\d{2}\b",
    "BR": r"\bBR-\d{2}\b",
    "NFR": r"\bNFR-\d{2}\b",
    "UC": r"\bUC-\d{2}\b",
    "AC": r"\bAC-\d{2}\b",
}
id_counts = {
    name: Counter(re.findall(pattern, "\n".join(texts)))
    for name, pattern in id_patterns.items()
}

terms = [
    "approver_employee_id", "leave_approval_logs", "leave_attachments",
    "leave_request_attachments", "remaining_days", "available_days",
    "pending_days", "used_days", "total_days", "ลาป่วยเกิน 3 วัน",
    "ทุก Action", "ทุก action", "ทุกครั้ง", "Future Scope", "Multi-step",
    "HR Approval", "Push Notification", "Real-time", "PDF Export",
]
term_locations = {}
for term in terms:
    matches = [record for record in records if term.casefold() in record.get("text", "").casefold()]
    term_locations[term] = matches

summary = {
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "records": len(records),
    "exact_duplicate_groups": len(exact_duplicates),
    "id_counts": {name: dict(sorted(counts.items())) for name, counts in id_counts.items()},
    "term_counts": {term: len(locations) for term, locations in term_locations.items()},
}

(OUT / "document_records.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
(OUT / "exact_duplicates.json").write_text(json.dumps(exact_duplicates, ensure_ascii=False, indent=2), encoding="utf-8")
(OUT / "term_locations.json").write_text(json.dumps(term_locations, ensure_ascii=False, indent=2), encoding="utf-8")
(OUT / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

with (OUT / "document.txt").open("w", encoding="utf-8") as f:
    for record in records:
        if record["kind"] == "paragraph":
            f.write(f"P{record['paragraph']:04d} [{record['style']}] {record['text']}\n")
        else:
            f.write(f"T{record['table']:02d}R{record['row']:02d} {record['text']}\n")

print(json.dumps(summary, ensure_ascii=False, indent=2))
