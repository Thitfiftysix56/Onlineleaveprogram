from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\User\Desktop\online-leavesystem")
DOCX = ROOT / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"
WORK = ROOT / ".codex-docx-audit"
BACKUP = WORK / "master-before-structural-refactor.docx"
REPORT = WORK / "structural-refactor-report.json"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def zip_hash(path: Path, name: str) -> str | None:
    with zipfile.ZipFile(path) as zf:
        if name not in zf.namelist():
            return None
        return hashlib.sha256(zf.read(name)).hexdigest()


def find_paragraph(doc: Document, text: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def first_rpr(paragraph: Paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def paragraph_node(template: Paragraph, text: str) -> object:
    node = deepcopy(template._p)
    for child in list(node):
        if child.tag != qn("w:pPr"):
            node.remove(child)
    run = OxmlElement("w:r")
    rpr = first_rpr(template)
    if rpr is not None:
        run.append(rpr)
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    node.append(run)
    return node


def clone_paragraph(paragraph: Paragraph, text: str | None = None) -> object:
    if text is None:
        return deepcopy(paragraph._p)
    return paragraph_node(paragraph, text)


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    new = paragraph_node(paragraph, text)
    paragraph._p.getparent().replace(paragraph._p, new)


def replace_cell_text(cell, text: str) -> None:
    template = cell.paragraphs[0]
    p = paragraph_node(template, text)
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(p)


def remove_range(start_node, end_exclusive) -> int:
    count = 0
    current = start_node
    while current is not None and current is not end_exclusive:
        nxt = current.getnext()
        current.getparent().remove(current)
        count += 1
        current = nxt
    if current is None and end_exclusive is not None:
        raise RuntimeError("Range end was not reached")
    return count


def remove_node(node) -> None:
    node.getparent().remove(node)


def insert_nodes_before(anchor_node, nodes: list[object]) -> None:
    for node in nodes:
        anchor_node.addprevious(node)


def heading_texts(doc: Document) -> list[str]:
    values = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and (
            text[0].isdigit()
            or text.startswith("FR-")
            or text.startswith("AC-")
            or text.startswith("UC-")
        ):
            values.append(text)
    return values


def main() -> None:
    WORK.mkdir(exist_ok=True)
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    shutil.copy2(DOCX, BACKUP)
    before_hash = sha256(BACKUP)
    before = Document(str(BACKUP))
    before_counts = {
        "paragraphs": len(before.paragraphs),
        "tables": len(before.tables),
        "sections": len(before.sections),
        "headings": heading_texts(before),
    }
    protected_parts = [
        "word/styles.xml",
        "word/header1.xml",
        "word/footer1.xml",
        "word/settings.xml",
        "word/theme/theme1.xml",
    ]
    protected_before = {name: zip_hash(BACKUP, name) for name in protected_parts}

    doc = before

    # Capture templates and Final Workflow content before removing duplicate sections.
    section_heading = find_paragraph(doc, "6. Workflow คำขอลา")
    subheading = find_paragraph(doc, "6.1 สถานะคำขอลา")
    nested_heading = find_paragraph(doc, "16.2.1 กรณี Save Draft")
    body_template = find_paragraph(doc, "ระบบใช้สถานะหลัก 5 สถานะสำหรับ Workflow การลาแบบ Supervisor อนุมัติขั้นเดียว")
    bullet_template = find_paragraph(doc, "Draft สามารถแก้ไขและลบได้โดยเจ้าของคำขอเท่านั้น")
    number_template = find_paragraph(doc, "Lock Leave Request และตรวจ status = pending")

    table_status_definitions = deepcopy(doc.tables[5]._tbl)
    table_validation = deepcopy(doc.tables[7]._tbl)
    table_cancel_rules = deepcopy(doc.tables[8]._tbl)
    table_final_transition = deepcopy(doc.tables[45]._tbl)
    nfr_table = doc.tables[11]
    leave_type_table = doc.tables[17]
    leave_requests_table = doc.tables[19]
    notification_table = doc.tables[23]
    uc19_table = doc.tables[69]

    approval_steps = [
        "Lock Leave Request และตรวจ status = pending",
        "ตรวจ Direct Team และป้องกัน Self Approval",
        "Lock Entitlement และตรวจ Balance ซ้ำ",
        "เพิ่ม used_days",
        "เปลี่ยน status เป็น approved",
        "บันทึก approver_employee_id และ approved_at",
        "สร้าง Notification ให้เจ้าของคำขอ",
        "บันทึก Audit Log action = leave_approved",
        "Commit โดยไม่ INSERT leave_approval_logs",
    ]
    rejection_steps = [
        "Lock Leave Request และตรวจสถานะ",
        "ตรวจ Direct Team และป้องกัน Self Decision",
        "ตรวจ Rejection Reason",
        "เปลี่ยน status เป็น rejected",
        "บันทึก rejection_reason, approver_employee_id และ rejected_at",
        "สร้าง Notification ให้เจ้าของคำขอ",
        "บันทึก Audit Log action = leave_rejected",
        "Commit โดยไม่ INSERT leave_approval_logs",
    ]
    state_rules = [
        "Draft สามารถแก้ไขและลบได้โดยเจ้าของคำขอเท่านั้น",
        "เมื่อ Submit แล้วสถานะเปลี่ยนเป็น Pending และแก้ไขข้อมูลคำขอไม่ได้",
        "Pending สามารถ Approve / Reject ได้เฉพาะ Supervisor ที่เป็น Direct Supervisor ของผู้ยื่นลาตาม employees.supervisor_id เท่านั้น",
        "Pending สามารถ Cancel โดยผู้ยื่นลาได้",
        "Approved, Rejected และ Cancelled เป็นสถานะสิ้นสุดในระบบเวอร์ชันปัจจุบัน",
        "Approved ต้องหักวันลาใน Transaction เดียวกับการเปลี่ยนสถานะ",
        "Rejected ต้องมีเหตุผลการปฏิเสธ",
        "ระบบบันทึก Event สำคัญที่กำหนดใน Audit Log ตามจุดที่ Backend รองรับ",
        "คำขอลาที่ Approved แล้วไม่สามารถแก้ไข ยกเลิก หรือลบได้ในระบบเวอร์ชันปัจจุบัน",
        "คำขอ Pending ต้องถูกนำไปคิดเป็น pending_days เพื่อป้องกันการส่งคำขอเกินสิทธิ์",
    ]

    section6_nodes = [
        paragraph_node(subheading, "6.1 ภาพรวม Workflow"),
        paragraph_node(body_template, "ผู้ใช้งานทุก Role สามารถสร้างคำขอลาของตนเองได้ตามสิทธิ์พื้นฐาน โดยระบบใช้ Supervisor Approval ขั้นเดียวและยึด employees.supervisor_id เป็น Source of Truth สำหรับ Direct Supervisor / Direct Team"),
    ]
    overview_steps = [
        "ผู้ใช้งานสร้าง Leave Request ของตนเอง",
        "บันทึกเป็น Draft หรือ Submit",
        "ระบบตรวจ Validation และ Business Rule ที่เกี่ยวข้อง",
        "ระบบสร้าง request_no เปลี่ยน status เป็น pending และแจ้ง Direct Supervisor",
        "Direct Supervisor ตรวจคำขอของ Direct Team โดยห้ามตัดสินคำขอตนเอง",
        "Supervisor อนุมัติหรือปฏิเสธ และระบบบันทึกผลการตัดสิน",
    ]
    for text in overview_steps:
        section6_nodes.append(paragraph_node(number_template, text))

    section6_nodes.extend([
        paragraph_node(subheading, "6.2 สถานะคำขอลา"),
        paragraph_node(body_template, "ระบบใช้สถานะหลัก 5 สถานะสำหรับ Workflow การลาแบบ Supervisor อนุมัติขั้นเดียว"),
        table_status_definitions,
        paragraph_node(subheading, "6.3 การสร้างและส่งคำขอลา"),
        paragraph_node(body_template, "ผู้ใช้งานทุก Role สามารถสร้าง Leave Request ของตนเอง เลือก Leave Type ระบุช่วงวันที่และ Reason พร้อมแนบเอกสารตาม Attachment Policy ของ Leave Type แล้วเลือก Save Draft หรือ Submit"),
        paragraph_node(nested_heading, "6.3.1 กรณี Save Draft"),
    ])
    for text in [
        "บันทึก status = draft",
        "เจ้าของแก้ไขหรือลบ Draft ได้",
        "เจ้าของจัดการ Attachment ได้",
        "สามารถ Submit ภายหลังได้",
    ]:
        section6_nodes.append(paragraph_node(bullet_template, text))
    section6_nodes.extend([
        paragraph_node(nested_heading, "6.3.2 กรณี Submit"),
        paragraph_node(body_template, "ระบบตรวจ Required Field, วันที่และปี, Reason 5–500 ตัวอักษร, Leave Type Active, Working Day, Minimum/Maximum Days, Entitlement, available_days ที่รวม pending_days, Overlap, Attachment Policy/Threshold, Direct Supervisor และสถานะ User/Employee"),
        paragraph_node(body_template, "เมื่อ Validation ผ่าน ระบบสร้าง request_no เปลี่ยน status เป็น pending ตรวจ Direct Supervisor จาก employees.supervisor_id และสร้าง Notification ให้ Supervisor โดยไม่บันทึก approver_employee_id ตอน Submit"),
        table_validation,
        paragraph_node(subheading, "6.4 Approval Workflow"),
        paragraph_node(body_template, "Supervisor อนุมัติได้เฉพาะคำขอ pending ของ Direct Team ตาม employees.supervisor_id และห้ามตัดสินคำขอของตนเอง ระบบตรวจ Permission, Entitlement และ Balance ซ้ำภายใน Transaction เพื่อป้องกัน Double Decision หรือ Concurrent Approval"),
    ])
    for text in approval_steps:
        section6_nodes.append(paragraph_node(number_template, text))
    section6_nodes.extend([
        paragraph_node(subheading, "6.5 Reject Workflow"),
        paragraph_node(body_template, "Supervisor ปฏิเสธได้เฉพาะคำขอ pending ของ Direct Team และต้องระบุ Rejection Reason"),
    ])
    for text in rejection_steps:
        section6_nodes.append(paragraph_node(number_template, text))
    section6_nodes.extend([
        paragraph_node(subheading, "6.6 Cancel Workflow"),
        paragraph_node(body_template, "เจ้าของคำขอสามารถ Cancel ได้เฉพาะ status = pending ระบบเปลี่ยนสถานะเป็น cancelled และบันทึก cancelled_at โดยไม่สร้าง Notification และไม่บันทึก Audit Log การ Cancel ไม่คืนสิทธิ์เพราะ pending_days เป็นค่าคำนวณและ used_days ยังไม่ถูกเพิ่ม"),
        table_cancel_rules,
        paragraph_node(subheading, "6.7 State Transition"),
        table_final_transition,
        paragraph_node(subheading, "6.8 State Rules"),
    ])
    for text in state_rules:
        section6_nodes.append(paragraph_node(bullet_template, text))
    section6_nodes.extend([
        paragraph_node(subheading, "6.9 Workflow การใช้งานแยกตาม Role"),
        paragraph_node(body_template, "ลำดับต่อไปนี้อธิบายทางเดินการใช้งานโดยสรุป ส่วนหน้าที่และ Permission หลักให้ยึดหัวข้อ 3–5"),
    ])
    role_flows = [
        "Employee: Login → Dashboard → สร้าง/ส่งคำขอลาของตนเอง → ติดตามคำขอและสิทธิ์การลา → จัดการการแจ้งเตือนและ Profile",
        "Supervisor: Login → Dashboard → ใช้ Own Leave Function หรือเปิดรายการ pending ของ Direct Team → Approve/Reject → ดู Team Report",
        "HR: Login → Dashboard → ใช้ Own Leave Function หรือจัดการข้อมูลพนักงาน ประเภทลา สิทธิ์ลา วันหยุด และรายงาน HR",
        "Admin: Login → Dashboard → ใช้ Own Leave Function หรือจัดการบัญชีผู้ใช้ แผนก ตำแหน่ง และดู Audit Log",
    ]
    for text in role_flows:
        section6_nodes.append(paragraph_node(bullet_template, text))

    # Replace the old Section 6 in place.
    heading7 = find_paragraph(doc, "7. Functional Requirements")
    removed_section6 = remove_range(section_heading._p.getnext(), heading7._p)
    insert_nodes_before(heading7._p, section6_nodes)

    # Future Scope is authoritative only at 2.2; remove the duplicate trailing list.
    duplicate_future = find_paragraph(doc, "16. Future Scope")
    use_case_section = find_paragraph(doc, "16. Workflow และ Use Case Final")
    page_break = use_case_section._p.getprevious()
    if not page_break.xpath(".//w:br[@w:type='page']"):
        raise RuntimeError("Expected page break before Section 16")
    removed_future = remove_range(duplicate_future._p, page_break)

    # Remove duplicated Workflow/Role content from the former Section 16.
    workflow_start = find_paragraph(doc, "16.1 Workflow ภาพรวมของระบบ")
    actor_heading = find_paragraph(doc, "16.8 Actor Definition")
    removed_duplicate_workflow = remove_range(workflow_start._p, actor_heading._p)

    # Retain only Use Case content and renumber it as a single Section 16.
    rename = {
        "16. Workflow และ Use Case Final": "16. Use Case",
        "16.8 Actor Definition": "16.1 Actor Definition",
        "16.9 Use Case Inventory": "16.2 Use Case Inventory",
        "16.10 Actor และ Use Case Mapping": "16.3 Actor และ Use Case Mapping",
        "16.11 Use Case Diagram Specification": "16.4 Use Case Diagram Specification",
        "16.12 Use Case Specification": "16.5 Use Case Specification",
        "16.13 Requirement และ Use Case Mapping": "16.6 Requirement และ Use Case Mapping",
    }
    for old, new in rename.items():
        replace_paragraph_text(find_paragraph(doc, old), new)

    # Remove the internal Final Consistency Check and PASS/FAIL table from the report.
    consistency_heading = find_paragraph(doc, "16.14 Final Consistency Check")
    sect_pr = doc.element.body.sectPr
    removed_internal_qa = remove_range(consistency_heading._p, sect_pr)

    # Consolidate Current/Future limitations and point to the single Future Scope source.
    replace_paragraph_text(
        find_paragraph(doc, "Multi-role User ยังไม่รองรับในระบบเวอร์ชันปัจจุบัน"),
        "Current Scope กำหนดให้ผู้ใช้งานหนึ่งคนมีหนึ่ง Role; Multi-role อยู่ในขอบเขตอนาคตตามหัวข้อ 2.2",
    )
    limitations_summary = find_paragraph(doc, "ระบบรองรับเฉพาะการลาเต็มวัน ยังไม่รองรับการลาครึ่งวัน")
    replace_paragraph_text(
        limitations_summary,
        "ข้อจำกัดและฟังก์ชันในอนาคต เช่น Half-day, Cross-year, Cancel After Approval, Multi-step/Delegate Approval, PDF Export, Advanced Analytics, Backup Assignee และ Multi-role สรุปไว้ที่หัวข้อ 2.2",
    )
    for text in [
        "ระบบไม่รองรับการยื่นคำขอลาข้ามปี",
        "ระบบไม่รองรับการยกเลิกคำขอที่ได้รับอนุมัติแล้ว",
        "ระบบไม่รองรับการอนุมัติหลายระดับ",
        "ระบบไม่รองรับการมอบหมายผู้อนุมัติแทน",
        "ระบบไม่รองรับ PDF Export",
        "Dashboard ปัจจุบันแสดงข้อมูลสรุปตาม Role แต่ยังไม่รองรับการวิเคราะห์ขั้นสูงหลายมิติ",
        "ระบบไม่รองรับ Backup Assignee",
        "ระบบไม่รองรับ Delegation",
        "ระบบไม่รองรับ Multi-role User",
    ]:
        remove_node(find_paragraph(doc, text)._p)
    replace_paragraph_text(
        find_paragraph(doc, "PDF เป็น Future Scope"),
        "PDF ไม่อยู่ใน Current Scope (ดูหัวข้อ 2.2)",
    )

    # Correct legacy/current documentation facts in place.
    replace_paragraph_text(
        find_paragraph(doc, "leave_requests 1 ──── * leave_attachments"),
        "leave_requests 1 ──── * leave_request_attachments",
    )
    replace_paragraph_text(
        find_paragraph(doc, "ระบบต้องตรวจสอบทั้ง file extension และ MIME type"),
        "ระบบต้องตรวจสอบ MIME type ตามรายการที่ Backend อนุญาต",
    )
    replace_paragraph_text(
        find_paragraph(doc, "กำหนดเงื่อนไขแนบเอกสาร เช่น ลาเกินกี่วันต้องแนบ"),
        "กำหนดเงื่อนไขแนบเอกสารตามเกณฑ์จำนวนวันที่เริ่มบังคับแนบ",
    )

    # Table correctness fixes.
    replace_cell_text(nfr_table.cell(12, 1), "จำกัดไฟล์แนบไม่เกิน 10 MB ต่อไฟล์ และตรวจ MIME type")
    replace_cell_text(leave_type_table.cell(5, 3), "เกณฑ์จำนวนวันที่เริ่มบังคับแนบเอกสาร")
    replace_cell_text(notification_table.cell(6, 3), "ประเภทการแจ้งเตือน เช่น leave-submitted, leave-approved, leave-rejected")

    # Add the production rejection_reason field to leave_requests without altering table style.
    if not any(row.cells[0].text.strip() == "rejection_reason" for row in leave_requests_table.rows):
        rejected_row = leave_requests_table.rows[13]._tr
        new_row = deepcopy(rejected_row)
        rejected_row.addnext(new_row)
        inserted = Table(leave_requests_table._tbl, leave_requests_table._parent).rows[14]
        values = ["rejection_reason", "TEXT", "NULL", "เหตุผลการปฏิเสธคำขอ"]
        for index, value in enumerate(values):
            replace_cell_text(inserted.cells[index], value)

    # Update the UC-19 current-scope reference without duplicating the Future Scope list.
    row = next(row for row in uc19_table.rows if row.cells[0].text.strip() == "Postconditions")
    replace_cell_text(row.cells[1], "ได้รายงานบนหน้าจอหรือไฟล์ Excel; PDF ไม่อยู่ใน Current Scope (ดูหัวข้อ 2.2)")

    doc.save(str(DOCX))

    after = Document(str(DOCX))
    headings_after = heading_texts(after)
    full_text = "\n".join(p.text for p in after.paragraphs) + "\n" + "\n".join(cell.text for table in after.tables for row in table.rows for cell in row.cells)
    protected_after = {name: zip_hash(DOCX, name) for name in protected_parts}
    use_case_headings = [p.text.strip() for p in after.paragraphs if p.text.strip().startswith("UC-")]
    section6_expected = [
        "6.1 ภาพรวม Workflow", "6.2 สถานะคำขอลา", "6.3 การสร้างและส่งคำขอลา",
        "6.3.1 กรณี Save Draft", "6.3.2 กรณี Submit", "6.4 Approval Workflow",
        "6.5 Reject Workflow", "6.6 Cancel Workflow", "6.7 State Transition",
        "6.8 State Rules", "6.9 Workflow การใช้งานแยกตาม Role",
    ]
    section16_expected = [
        "16. Use Case", "16.1 Actor Definition", "16.2 Use Case Inventory",
        "16.3 Actor และ Use Case Mapping", "16.4 Use Case Diagram Specification",
        "16.5 Use Case Specification", "16.6 Requirement และ Use Case Mapping",
    ]
    report = {
        "document": str(DOCX),
        "sha256_before": before_hash,
        "sha256_after": sha256(DOCX),
        "before": before_counts,
        "after": {
            "paragraphs": len(after.paragraphs),
            "tables": len(after.tables),
            "sections": len(after.sections),
            "headings": headings_after,
        },
        "removed_element_counts": {
            "old_section6": removed_section6,
            "duplicate_future_scope": removed_future,
            "duplicate_workflow_role": removed_duplicate_workflow,
            "internal_consistency_qa": removed_internal_qa,
        },
        "section6_complete": all(value in headings_after for value in section6_expected),
        "section16_complete": all(value in headings_after for value in section16_expected),
        "duplicate_section16_heading_absent": "16. Future Scope" not in headings_after and "16. Workflow และ Use Case Final" not in headings_after,
        "final_consistency_heading_absent": "16.14 Final Consistency Check" not in headings_after,
        "use_case_count": len(use_case_headings),
        "legacy_leave_attachments_absent": "leave_attachments" not in full_text,
        "current_attachment_name_present": "leave_request_attachments" in full_text,
        "rejection_reason_present": "rejection_reason" in full_text,
        "extension_claim_absent": "file extension" not in full_text and "MIME type / extension" not in full_text,
        "single_future_scope_heading": headings_after.count("2.2 ขอบเขตในอนาคต (Future Scope)") == 1,
        "protected_parts_unchanged": {name: protected_before[name] == protected_after[name] for name in protected_parts},
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))

    failures = []
    for key in [
        "section6_complete", "section16_complete", "duplicate_section16_heading_absent",
        "final_consistency_heading_absent", "legacy_leave_attachments_absent",
        "current_attachment_name_present", "rejection_reason_present", "extension_claim_absent",
        "single_future_scope_heading",
    ]:
        if not report[key]:
            failures.append(key)
    if report["use_case_count"] != 24:
        failures.append("use_case_count")
    if not all(report["protected_parts_unchanged"].values()):
        failures.append("protected_parts_unchanged")
    if failures:
        raise RuntimeError("Structural QA failed: " + ", ".join(failures))


if __name__ == "__main__":
    main()
