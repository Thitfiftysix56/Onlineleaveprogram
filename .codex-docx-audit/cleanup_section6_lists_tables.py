from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\User\Desktop\online-leavesystem")
DOCX = ROOT / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"


def find_paragraph(doc, text):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def new_number_instance(doc, source_paragraph):
    numbering = doc.part.numbering_part.element
    old_num_nodes = source_paragraph._p.xpath("./w:pPr/w:numPr/w:numId")
    if not old_num_nodes:
        raise RuntimeError(f"Paragraph is not numbered: {source_paragraph.text}")
    old_num_id = old_num_nodes[0].get(qn("w:val"))
    old_num = next(x for x in numbering.findall(qn("w:num")) if x.get(qn("w:numId")) == old_num_id)
    abstract_id = old_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    current_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_id = max(current_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_num_id(paragraph, num_id):
    node = paragraph._p.xpath("./w:pPr/w:numPr/w:numId")[0]
    node.set(qn("w:val"), str(num_id))


doc = Document(str(DOCX))

# The first copies are in Section 6; later tables remain the authoritative FR tables.
for index in sorted([7, 6], reverse=True):
    table = doc.tables[index]
    header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
    expected = "Current Status | Action | Result" if index == 7 else "Validation | เงื่อนไข"
    if header != expected:
        raise RuntimeError(f"Unexpected table {index}: {header}")
    table._tbl.getparent().remove(table._tbl)

groups = [
    ("6.1 ภาพรวม Workflow", "6.2 สถานะคำขอลา", [
        "ผู้ใช้งานสร้าง Leave Request ของตนเอง",
        "บันทึกเป็น Draft หรือ Submit",
        "ระบบตรวจ Validation และ Business Rule ที่เกี่ยวข้อง",
        "ระบบสร้าง request_no เปลี่ยน status เป็น pending และแจ้ง Direct Supervisor",
        "Direct Supervisor ตรวจคำขอของ Direct Team โดยห้ามตัดสินคำขอตนเอง",
        "Supervisor อนุมัติหรือปฏิเสธ และระบบบันทึกผลการตัดสิน",
    ]),
    ("6.4 Approval Workflow", "6.5 Reject Workflow", [
        "Lock Leave Request และตรวจ status = pending",
        "ตรวจ Direct Team และป้องกัน Self Approval",
        "Lock Entitlement และตรวจ Balance ซ้ำ",
        "เพิ่ม used_days",
        "เปลี่ยน status เป็น approved",
        "บันทึก approver_employee_id และ approved_at",
        "สร้าง Notification ให้เจ้าของคำขอ",
        "บันทึก Audit Log action = leave_approved",
        "Commit โดยไม่ INSERT leave_approval_logs",
    ]),
    ("6.5 Reject Workflow", "6.6 Cancel Workflow", [
        "Lock Leave Request และตรวจสถานะ",
        "ตรวจ Direct Team และป้องกัน Self Decision",
        "ตรวจ Rejection Reason",
        "เปลี่ยน status เป็น rejected",
        "บันทึก rejection_reason, approver_employee_id และ rejected_at",
        "สร้าง Notification ให้เจ้าของคำขอ",
        "บันทึก Audit Log action = leave_rejected",
        "Commit โดยไม่ INSERT leave_approval_logs",
    ]),
]

for start_text, end_text, texts in groups:
    all_paragraphs = doc.paragraphs
    start_matches = [i for i, p in enumerate(all_paragraphs) if p.text.strip() == start_text]
    end_matches = [i for i, p in enumerate(all_paragraphs) if p.text.strip() == end_text]
    if len(start_matches) != 1 or len(end_matches) != 1:
        raise RuntimeError(
            f"Expected one heading pair {start_text!r} -> {end_text!r}, "
            f"found {len(start_matches)} and {len(end_matches)}"
        )
    start_index = start_matches[0]
    end_index = end_matches[0]
    scope = all_paragraphs[start_index + 1:end_index]
    paragraphs = []
    for text in texts:
        matches = [p for p in scope if p.text.strip() == text]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one scoped paragraph {text!r}, found {len(matches)}")
        paragraphs.append(matches[0])
    num_id = new_number_instance(doc, paragraphs[0])
    for paragraph in paragraphs:
        apply_num_id(paragraph, num_id)

doc.save(str(DOCX))
print({"tables": len(Document(str(DOCX)).tables), "numbered_groups_reset": len(groups)})
