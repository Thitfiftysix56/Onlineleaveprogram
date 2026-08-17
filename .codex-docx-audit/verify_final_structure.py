import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\User\Desktop\online-leavesystem")
DOCX = ROOT / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"


def paragraph_num_id(paragraph):
    nodes = paragraph._p.xpath("./w:pPr/w:numPr/w:numId")
    return nodes[0].get(qn("w:val")) if nodes else None


def numbering_start(doc, num_id, level="0"):
    numbering = doc.part.numbering_part.element
    num = next(n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == num_id)
    for override in num.findall(qn("w:lvlOverride")):
        if override.get(qn("w:ilvl")) == level:
            start = override.find(qn("w:startOverride"))
            if start is not None:
                return int(start.get(qn("w:val")))
    abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next(
        n for n in numbering.findall(qn("w:abstractNum"))
        if n.get(qn("w:abstractNumId")) == abstract_id
    )
    level_node = next(
        n for n in abstract.findall(qn("w:lvl"))
        if n.get(qn("w:ilvl")) == level
    )
    start = level_node.find(qn("w:start"))
    return int(start.get(qn("w:val"))) if start is not None else 1


doc = Document(str(DOCX))
all_text = "\n".join(p.text for p in doc.paragraphs)
all_text += "\n" + "\n".join(
    " | ".join(c.text for c in row.cells)
    for table in doc.tables for row in table.rows
)

headings = [p.text.strip() for p in doc.paragraphs if re.match(r"^(?:[1-9]|1[0-6])(?:\.\d+)*\.?(?:\s|$)", p.text.strip())]
heading_numbers = {
    match.group(1)
    for heading in headings
    if (match := re.match(r"^((?:[1-9]|1[0-6])(?:\.\d+)*)\.?(?:\s|$)", heading))
}
referenced_heading_numbers = sorted(set(re.findall(r"หัวข้อ\s+((?:[1-9]|1[0-6])(?:\.\d+)*)", all_text)))
broken_heading_references = [number for number in referenced_heading_numbers if number not in heading_numbers]
if broken_heading_references:
    raise RuntimeError(f"Broken heading references: {broken_heading_references}")

uc_results = []
for number in range(1, 25):
    uc = f"UC-{number:02d}"
    candidates = [
        t for t in doc.tables
        if any(
            len(row.cells) >= 2
            and row.cells[0].text.strip() == "Use Case ID"
            and row.cells[1].text.strip() == uc
            for row in t.rows
        )
    ]
    if len(candidates) != 1:
        raise RuntimeError(f"Expected one specification table for {uc}, found {len(candidates)}")
    table = candidates[0]
    flow_rows = [r for r in table.rows if r.cells[0].text.strip() == "Main Flow"]
    if len(flow_rows) != 1:
        raise RuntimeError(f"Expected one Main Flow row for {uc}, found {len(flow_rows)}")
    paragraphs = [p for p in flow_rows[0].cells[1].paragraphs if p.text.strip()]
    if not paragraphs:
        raise RuntimeError(f"Empty Main Flow for {uc}")
    num_id = paragraph_num_id(paragraphs[0])
    if num_id is None:
        raise RuntimeError(f"First Main Flow step is not numbered for {uc}")
    starts_at = numbering_start(doc, num_id)
    uc_results.append({"uc": uc, "steps": len(paragraphs), "num_id": num_id, "starts_at": starts_at})

for prefix, maximum in (("FR", 21), ("BR", 40), ("NFR", 16), ("AC", 10), ("UC", 24)):
    missing = [f"{prefix}-{i:02d}" for i in range(1, maximum + 1) if f"{prefix}-{i:02d}" not in all_text]
    if missing:
        raise RuntimeError(f"Missing identifiers for {prefix}: {missing}")

for forbidden in (
    "16. Future Scope",
    "Workflow และ Use Case Final",
    "16.14 Final Consistency Check",
    "leave_attachments",
):
    if forbidden in all_text:
        raise RuntimeError(f"Forbidden legacy text remains: {forbidden}")

section6 = [h for h in headings if h.startswith("6.")]
section16 = [h for h in headings if h.startswith("16.")]
result = {
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "section6_headings": section6,
    "section16_headings": section16,
    "referenced_heading_numbers": referenced_heading_numbers,
    "broken_heading_references": broken_heading_references,
    "use_case_count": len(uc_results),
    "main_flow_all_start_at_1": all(x["starts_at"] == 1 for x in uc_results),
    "main_flow_distinct_num_ids": len({x["num_id"] for x in uc_results}),
    "use_cases": uc_results,
}

out = ROOT / ".codex-docx-audit" / "final-structure-verification.json"
out.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
print(json.dumps(result, ensure_ascii=False, indent=2))
