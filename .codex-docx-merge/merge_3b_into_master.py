from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\User\Desktop\online-leavesystem")
OUT = ROOT / "Final Requirement Output"
MASTER = OUT / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"
SOURCE = OUT / "ระบบอนุมัติใบลาออนไลน์ Final Workflow and Use Case.docx"
WORK = ROOT / ".codex-docx-merge"
BACKUP = WORK / "master-before-3b-merge.docx"
REPORT = WORK / "merge-qa.json"


MAJOR_HEADINGS = {
    "1. Workflow ภาพรวมของระบบ": "16.1 Workflow ภาพรวมของระบบ",
    "2. Leave Request Workflow": "16.2 Leave Request Workflow",
    "3. Approval Workflow": "16.3 Approval Workflow",
    "4. Reject Workflow": "16.4 Reject Workflow",
    "5. Cancel Workflow": "16.5 Cancel Workflow",
    "6. Status Transition Table": "16.6 Status Transition",
    "7. Workflow แยกตาม Role": "16.7 Workflow แยกตาม Role",
    "8. Actor Definition": "16.8 Actor Definition",
    "9. Use Case Inventory": "16.9 Use Case Inventory",
    "10. Actor ↔ Use Case Mapping": "16.10 Actor และ Use Case Mapping",
    "11. Use Case Diagram Specification": "16.11 Use Case Diagram Specification",
    "12. Use Case Specification": "16.12 Use Case Specification",
    "13. Requirement ↔ Use Case Mapping": "16.13 Requirement และ Use Case Mapping",
    "14. Final Consistency Check": "16.14 Final Consistency Check",
}

SUBHEADINGS = {
    "2.1 กรณี Save Draft": "16.2.1 กรณี Save Draft",
    "2.2 กรณี Submit": "16.2.2 กรณี Submit",
}

TABLE_WIDTHS = {
    0: [1.20, 1.12, 1.42, 1.46, 1.85],
    1: [1.10, 2.30, 3.65],
    2: [1.15, 2.65, 3.25],
    3: [0.72, 2.18, 1.45, 2.70],
    4: [2.08, 0.82, 0.92, 0.72, 0.78, 1.73],
    5: [1.25, 5.80],
    30: [0.67, 1.65, 0.78, 1.52, 1.12, 1.31],
    31: [2.25, 3.85, 0.95],
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def zip_part_hash(path: Path, name: str) -> str | None:
    with zipfile.ZipFile(path) as zf:
        if name not in zf.namelist():
            return None
        return hashlib.sha256(zf.read(name)).hexdigest()


def body_fingerprint(doc: Document) -> list[str]:
    out = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        out.append(hashlib.sha256(child.xml.encode("utf-8")).hexdigest())
    return out


def table_text_fingerprint(doc: Document) -> list[list[list[str]]]:
    return [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]


def get_app_pages(path: Path) -> int | None:
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("docProps/app.xml")
        from lxml import etree

        root = etree.fromstring(xml)
        nodes = root.xpath("//*[local-name()='Pages']")
        return int(nodes[0].text) if nodes and nodes[0].text else None
    except Exception:
        return None


def replace_ppr(dst: Paragraph, src: Paragraph) -> None:
    if dst._p.pPr is not None:
        dst._p.remove(dst._p.pPr)
    if src._p.pPr is not None:
        dst._p.insert(0, deepcopy(src._p.pPr))


def get_first_rpr(paragraph: Paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def add_formatted_paragraph(doc: Document, text: str, template: Paragraph) -> Paragraph:
    p = doc.add_paragraph()
    replace_ppr(p, template)
    run = p.add_run(text)
    rpr = get_first_rpr(template)
    if rpr is not None:
        run._r.insert(0, rpr)
    return p


def next_numbering_ids(doc: Document) -> tuple[int, int]:
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc: Document, bullet: bool = False) -> int:
    abstract_id, num_id = next_numbering_ids(doc)
    root = doc.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "240")
    ppr.append(ind)
    lvl.append(ppr)
    # Do not force Symbol: the Master document's TH Sarabun New renders the
    # Unicode bullet correctly and avoids missing-glyph boxes in PDF export.
    abstract.append(lvl)
    root.insert(len(root.findall(qn("w:abstractNum"))), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    root.append(num)
    return num_id


def apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:numPr"))
    if old is not None:
        ppr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    ppr.append(num_pr)


def add_numbered_paragraph(doc: Document, text: str, template: Paragraph, num_id: int) -> Paragraph:
    p = add_formatted_paragraph(doc, text, template)
    apply_numbering(p, num_id)
    return p


def copy_cell_base_format(src: _Cell, dst: _Cell) -> None:
    tc_pr = src._tc.tcPr
    if tc_pr is not None:
        for child in list(dst._tc.tcPr):
            dst._tc.tcPr.remove(child)
        for child in tc_pr:
            if child.tag != qn("w:tcW"):
                dst._tc.tcPr.append(deepcopy(child))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_indent(table: Table, twips: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblInd"))
    if old is not None:
        tbl_pr.remove(old)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(twips))
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def normalize_twips(widths_inches: list[float], total_twips: int = 9350) -> list[int]:
    raw = [value * 1440 for value in widths_inches]
    raw_total = sum(raw)
    result = [round(value * total_twips / raw_total) for value in raw]
    result[-1] += total_twips - sum(result)
    return result


def set_table_geometry(table: Table, widths_inches: list[float], total_twips: int = 9350) -> None:
    widths = normalize_twips(widths_inches, total_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.tcPr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.tcPr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell: _Cell) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)
    cell._tc.append(OxmlElement("w:p"))


def format_cell_paragraph(p: Paragraph, template: Paragraph, text: str, bold: bool | None = None) -> None:
    replace_ppr(p, template)
    run = p.add_run(text)
    rpr = get_first_rpr(template)
    if rpr is not None:
        run._r.insert(0, rpr)
    if bold is not None:
        run.bold = bold


def add_cell_paragraph(cell: _Cell, template: Paragraph, text: str, num_id: int | None = None) -> Paragraph:
    p = cell.add_paragraph()
    format_cell_paragraph(p, template, text)
    if num_id is not None:
        apply_numbering(p, num_id)
    return p


def source_cell_lines(cell: _Cell) -> list[tuple[str, str]]:
    lines = []
    for p in cell.paragraphs:
        text = p.text.strip()
        if text:
            lines.append((p.style.name if p.style is not None else "Normal", text))
    return lines


def build_table(
    doc: Document,
    source_table: Table,
    table_index: int,
    master_header_cell: _Cell,
    master_body_cell: _Cell,
    header_template: Paragraph,
    body_template: Paragraph,
) -> Table:
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    set_table_indent(table)
    widths = TABLE_WIDTHS[table_index]

    for r_idx, row in enumerate(table.rows):
        set_row_no_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            copy_cell_base_format(master_header_cell if r_idx == 0 else master_body_cell, cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            clear_cell(cell)
            lines = source_cell_lines(source_table.cell(r_idx, c_idx))
            if not lines:
                lines = [("Normal", "")]
            p = cell.paragraphs[0]
            format_cell_paragraph(
                p,
                header_template if r_idx == 0 else body_template,
                lines[0][1],
                True if r_idx == 0 else None,
            )
            for _, text in lines[1:]:
                add_cell_paragraph(cell, body_template, text)
    set_table_geometry(table, widths)
    return table


def build_uc_table(
    doc: Document,
    source_table: Table,
    master_header_cell: _Cell,
    master_body_cell: _Cell,
    header_template: Paragraph,
    body_template: Paragraph,
    bullet_num_id: int,
) -> Table:
    table = doc.add_table(rows=len(source_table.rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_indent(table)
    widths = [1.65, 5.40]
    main_flow_num_id = create_numbering(doc, bullet=False)

    for r_idx, row in enumerate(table.rows):
        set_row_no_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        label = source_table.cell(r_idx, 0).text.strip()
        for c_idx, cell in enumerate(row.cells):
            copy_cell_base_format(master_header_cell if r_idx == 0 else master_body_cell, cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            clear_cell(cell)
            lines = source_cell_lines(source_table.cell(r_idx, c_idx))
            if not lines:
                lines = [("Normal", "")]
            p = cell.paragraphs[0]
            format_cell_paragraph(
                p,
                header_template if r_idx == 0 else body_template,
                lines[0][1],
                True if r_idx == 0 else None,
            )
            if c_idx == 1 and r_idx > 0:
                if label == "Main Flow":
                    apply_numbering(p, main_flow_num_id)
                elif label in {"Alternative Flow", "Exception Flow"}:
                    apply_numbering(p, bullet_num_id)
            for _, text in lines[1:]:
                num_id = None
                if c_idx == 1 and label == "Main Flow":
                    num_id = main_flow_num_id
                elif c_idx == 1 and label in {"Alternative Flow", "Exception Flow"}:
                    num_id = bullet_num_id
                add_cell_paragraph(cell, body_template, text, num_id=num_id)
    set_table_geometry(table, widths)
    return table


def has_page_break(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def get_text_from_child(child, parent) -> str:
    return Paragraph(child, parent).text.strip()


def main() -> None:
    WORK.mkdir(parents=True, exist_ok=True)
    if not MASTER.exists() or not SOURCE.exists():
        raise FileNotFoundError("Master or source document is missing")

    master_hash_before = sha256(MASTER)
    shutil.copy2(MASTER, BACKUP)
    before = Document(str(BACKUP))
    source = Document(str(SOURCE))

    if any(p.text.strip().startswith("16. Workflow และ Use Case Final") for p in before.paragraphs):
        raise RuntimeError("Master already contains Section 16; merge aborted to avoid duplication")

    before_paragraph_texts = [p.text for p in before.paragraphs]
    before_tables = table_text_fingerprint(before)
    before_table_count = len(before.tables)
    before_paragraph_count = len(before.paragraphs)
    before_body_fp = body_fingerprint(before)
    before_sections = [
        {
            "width": s.page_width,
            "height": s.page_height,
            "top": s.top_margin,
            "bottom": s.bottom_margin,
            "left": s.left_margin,
            "right": s.right_margin,
            "orientation": str(s.orientation),
        }
        for s in before.sections
    ]
    protected_parts = [
        "word/styles.xml",
        "word/header1.xml",
        "word/footer1.xml",
        "word/settings.xml",
        "word/theme/theme1.xml",
    ]
    protected_before = {part: zip_part_hash(BACKUP, part) for part in protected_parts}

    doc = before
    section_heading_template = doc.paragraphs[2]
    subheading_template = doc.paragraphs[26]
    body_template = doc.paragraphs[3]
    master_header_cell = doc.tables[0].cell(0, 0)
    master_body_cell = doc.tables[0].cell(1, 0)
    header_template = master_header_cell.paragraphs[0]
    table_body_template = master_body_cell.paragraphs[0]

    # Exactly one explicit page break before Section 16.
    pbreak = doc.add_paragraph()
    pbreak.add_run().add_break(WD_BREAK.PAGE)
    add_formatted_paragraph(doc, "16. Workflow และ Use Case Final", section_heading_template)

    bullet_num_id = create_numbering(doc, bullet=True)
    current_number_num_id = None
    previous_was_number = False
    source_table_index = 0
    started = False

    for child in source.element.body.iterchildren():
        tag = child.tag
        if tag == qn("w:sectPr"):
            continue
        if tag == qn("w:p"):
            p_src = Paragraph(child, source._body)
            text = p_src.text.strip()
            if text == "1. Workflow ภาพรวมของระบบ":
                started = True
            if not started:
                continue
            if not text or has_page_break(p_src):
                previous_was_number = False
                continue
            if text in MAJOR_HEADINGS:
                add_formatted_paragraph(doc, MAJOR_HEADINGS[text], section_heading_template)
                previous_was_number = False
            elif text in SUBHEADINGS:
                add_formatted_paragraph(doc, SUBHEADINGS[text], subheading_template)
                previous_was_number = False
            elif p_src.style.name == "Heading 2":
                add_formatted_paragraph(doc, text, subheading_template)
                previous_was_number = False
            elif p_src.style.name == "List Number":
                if not previous_was_number:
                    current_number_num_id = create_numbering(doc, bullet=False)
                add_numbered_paragraph(doc, text, body_template, current_number_num_id)
                previous_was_number = True
            elif p_src.style.name == "List Bullet":
                add_numbered_paragraph(doc, text, body_template, bullet_num_id)
                previous_was_number = False
            else:
                add_formatted_paragraph(doc, text, body_template)
                previous_was_number = False
        elif tag == qn("w:tbl"):
            if not started:
                source_table_index += 1
                continue
            table_src = Table(child, source._body)
            if 6 <= source_table_index <= 29:
                build_uc_table(
                    doc,
                    table_src,
                    master_header_cell,
                    master_body_cell,
                    header_template,
                    table_body_template,
                    bullet_num_id,
                )
            else:
                build_table(
                    doc,
                    table_src,
                    source_table_index,
                    master_header_cell,
                    master_body_cell,
                    header_template,
                    table_body_template,
                )
            source_table_index += 1
            previous_was_number = False

    doc.save(str(MASTER))

    after = Document(str(MASTER))
    after_sections = [
        {
            "width": s.page_width,
            "height": s.page_height,
            "top": s.top_margin,
            "bottom": s.bottom_margin,
            "left": s.left_margin,
            "right": s.right_margin,
            "orientation": str(s.orientation),
        }
        for s in after.sections
    ]
    after_body_fp = body_fingerprint(after)
    protected_after = {part: zip_part_hash(MASTER, part) for part in protected_parts}

    # Main Flow QA: every UC table must have a distinct numId and visible first value = 1.
    appended_tables = after.tables[before_table_count:]
    uc_tables = [t for t in appended_tables if t.cell(0, 0).text.strip() == "หัวข้อ" and any(r.cells[0].text.strip() == "Main Flow" for r in t.rows)]
    main_flow_num_ids = []
    main_flow_counts = []
    for table in uc_tables:
        row = next(r for r in table.rows if r.cells[0].text.strip() == "Main Flow")
        paragraphs = [p for p in row.cells[1].paragraphs if p.text.strip()]
        ids = []
        for p in paragraphs:
            nodes = p._p.xpath("./w:pPr/w:numPr/w:numId")
            ids.append(nodes[0].get(qn("w:val")) if nodes else None)
        main_flow_num_ids.append(ids[0] if ids else None)
        main_flow_counts.append(len(paragraphs))

    section16_texts = [p.text.strip() for p in after.paragraphs[before_paragraph_count:] if p.text.strip()]
    use_case_headings = [t for t in section16_texts if t.startswith("UC-")]
    fr_table = appended_tables[-2]
    fr_ids = [fr_table.cell(i, 0).text.strip() for i in range(1, len(fr_table.rows))]

    before_prefix_preserved = after_body_fp[: len(before_body_fp)] == before_body_fp
    report = {
        "master": str(MASTER),
        "source": str(SOURCE),
        "backup": str(BACKUP),
        "sha256_before": master_hash_before,
        "sha256_after": sha256(MASTER),
        "app_pages_before": get_app_pages(BACKUP),
        "app_pages_after_save": get_app_pages(MASTER),
        "paragraphs_before": before_paragraph_count,
        "paragraphs_after": len(after.paragraphs),
        "tables_before": before_table_count,
        "tables_after": len(after.tables),
        "sections_before": len(before.sections),
        "sections_after": len(after.sections),
        "section_settings_unchanged": before_sections == after_sections,
        "paragraph_text_prefix_unchanged": [p.text for p in after.paragraphs[: len(before_paragraph_texts)]] == before_paragraph_texts,
        "table_text_prefix_unchanged": table_text_fingerprint(after)[: len(before_tables)] == before_tables,
        "body_xml_prefix_unchanged": before_prefix_preserved,
        "protected_parts_unchanged": {part: protected_before[part] == protected_after[part] for part in protected_parts},
        "section16_headings_present": {value: value in section16_texts for value in MAJOR_HEADINGS.values()},
        "use_case_heading_count": len(use_case_headings),
        "uc_table_count": len(uc_tables),
        "main_flow_distinct_num_ids": len(set(main_flow_num_ids)) == 24,
        "main_flow_all_numbered": all(x is not None for x in main_flow_num_ids) and len(main_flow_num_ids) == 24,
        "main_flow_step_counts": main_flow_counts,
        "fr_ids": fr_ids,
        "fr01_to_fr21_complete": fr_ids == [f"FR-{i:02d}" for i in range(1, 22)],
        "page_breaks_in_appended_top_level": sum(len(p._p.xpath(".//w:br[@w:type='page']")) for p in after.paragraphs[before_paragraph_count:]),
        "source_table_count_consumed": source_table_index,
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))

    failures = []
    for key in [
        "section_settings_unchanged",
        "paragraph_text_prefix_unchanged",
        "table_text_prefix_unchanged",
        "body_xml_prefix_unchanged",
        "main_flow_distinct_num_ids",
        "main_flow_all_numbered",
        "fr01_to_fr21_complete",
    ]:
        if not report[key]:
            failures.append(key)
    if not all(report["protected_parts_unchanged"].values()):
        failures.append("protected_parts_unchanged")
    if not all(report["section16_headings_present"].values()):
        failures.append("section16_headings_present")
    if report["use_case_heading_count"] != 24 or report["uc_table_count"] != 24:
        failures.append("use_case_count")
    if report["page_breaks_in_appended_top_level"] != 1:
        failures.append("single_page_break")
    if failures:
        raise RuntimeError("QA failed: " + ", ".join(failures))


if __name__ == "__main__":
    main()
