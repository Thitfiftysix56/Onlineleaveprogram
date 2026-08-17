from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


root = Path(r"C:\Users\User\Desktop\online-leavesystem")
current = root / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx"
recovered = root / ".codex-docx-merge" / "master-1-15-recovered.docx"

doc = Document(str(current))
body = doc.element.body
children = list(body.iterchildren())
section16_index = None
for index, child in enumerate(children):
    if child.tag == qn("w:p"):
        if Paragraph(child, doc._body).text.strip() == "16. Workflow และ Use Case Final":
            section16_index = index
            break
if section16_index is None:
    raise RuntimeError("Section 16 heading not found")

start = section16_index
if start > 0 and children[start - 1].tag == qn("w:p"):
    previous = Paragraph(children[start - 1], doc._body)
    if previous._p.xpath(".//w:br[@w:type='page']"):
        start -= 1

for child in children[start:]:
    if child.tag != qn("w:sectPr"):
        body.remove(child)

doc.save(str(recovered))
check = Document(str(recovered))
assert len(check.paragraphs) == 890
assert len(check.tables) == 45
assert not any(p.text.strip().startswith("16. Workflow") for p in check.paragraphs)
print(recovered)
