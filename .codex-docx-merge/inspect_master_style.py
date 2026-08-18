from pathlib import Path
import sys
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
root = Path(__file__).resolve().parents[1]
master = Document(root / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Requirement.docx")
source = Document(root / "Final Requirement Output" / "ระบบอนุมัติใบลาออนไลน์ Final Workflow and Use Case.docx")

print("MASTER", len(master.paragraphs), len(master.tables), len(master.sections))
print("SOURCE", len(source.paragraphs), len(source.tables), len(source.sections))
for i in [2, 25, 26, 169, 813]:
    p = master.paragraphs[i]
    print("P", i, repr(p.text), p.style.name, p.alignment, p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.line_spacing)
    print([(r.font.name, r.font.size.pt if r.font.size else None, r.bold, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None) for r in p.runs[:4]])

for ti in [0, 4, 10, 42]:
    t = master.tables[ti]
    print("T", ti, len(t.rows), len(t.columns), t.style.name if t.style else None)
    for ri in range(min(2, len(t.rows))):
        print(" row", ri, [c.text for c in t.rows[ri].cells])
        for c in t.rows[ri].cells[:2]:
            shd = c._tc.get_or_add_tcPr().find(qn("w:shd"))
            print("  shd", shd.get(qn("w:fill")) if shd is not None else None, [(r.font.name, r.font.size.pt if r.font.size else None, r.bold) for r in c.paragraphs[0].runs])

print("HEADERS", [p.text for p in master.sections[0].header.paragraphs])
print("FOOTERS", [p.text for p in master.sections[0].footer.paragraphs])
